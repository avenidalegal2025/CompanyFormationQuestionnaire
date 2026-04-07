import os
import json
import tempfile
import boto3
import re
import base64
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Constants
TEMPLATE_BUCKET = os.environ.get('TEMPLATE_BUCKET', 'avenida-legal-documents')
OUTPUT_BUCKET = os.environ.get('OUTPUT_BUCKET', 'avenida-legal-documents')
BUCKET_NAME = os.environ.get('BUCKET_NAME', TEMPLATE_BUCKET)

# Initialize S3 client
s3_client = boto3.client('s3', region_name=os.environ.get('AWS_REGION', 'us-west-1'))


def download_from_s3(bucket, key, local_path):
    """Download file from S3 to local path"""
    print(f"===> Downloading s3://{bucket}/{key} to {local_path}")
    s3_client.download_file(bucket, key, local_path)
    print("===> Downloaded successfully")


def upload_to_s3(local_path, bucket, key):
    """Upload file from local path to S3"""
    print(f"===> Uploading {local_path} to s3://{bucket}/{key}")
    s3_client.upload_file(local_path, bucket, key)
    print("===> Uploaded successfully")


def extract_s3_info(url):
    """Extract bucket and key from S3 URL"""
    if not url:
        return None, None

    if url.startswith('s3://'):
        parts = url[5:].split('/', 1)
        if len(parts) == 2:
            return parts[0], parts[1]
        return parts[0], ''

    if url.startswith('https://'):
        match = re.match(r'https://([^.]+)\.s3[^/]*\.amazonaws\.com/(.+)', url)
        if match:
            return match.group(1), match.group(2)

        match = re.match(r'https://s3[^/]*\.amazonaws\.com/([^/]+)/(.+)', url)
        if match:
            return match.group(1), match.group(2)

    return None, None


# ---------- Run-level replacement (non-destructive) ----------

def _insert_run_after(paragraph, run, text):
    """Insert a new run after the given run, returning the new run."""
    new_run = paragraph.add_run(text)
    run._element.addnext(new_run._element)
    return new_run


def _apply_run_format(source_run, target_run):
    """Copy formatting from source_run to target_run."""
    if source_run._element.rPr is not None:
        target_run._element.insert(0, deepcopy(source_run._element.rPr))


def _enforce_font(run):
    """Ensure a modified run uses 12pt Times New Roman."""
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def _replace_span_in_paragraph(paragraph, start, end, value):
    """Replace a character span [start, end) in the paragraph's runs with value.
    Preserves run formatting. Returns the run containing the replacement value, or None."""
    pos = 0
    run_start_idx = None
    run_end_idx = None
    run_start_offset = 0
    run_end_offset = 0

    for i, run in enumerate(paragraph.runs):
        run_len = len(run.text)
        if run_start_idx is None and pos + run_len > start:
            run_start_idx = i
            run_start_offset = start - pos
        if run_start_idx is not None and pos + run_len >= end:
            run_end_idx = i
            run_end_offset = end - pos
            break
        pos += run_len

    if run_start_idx is None or run_end_idx is None:
        return None

    if run_start_idx == run_end_idx:
        run = paragraph.runs[run_start_idx]
        prefix = run.text[:run_start_offset]
        suffix = run.text[run_end_offset:]

        if prefix:
            run.text = prefix
            value_run = _insert_run_after(paragraph, run, value)
            _apply_run_format(run, value_run)
        else:
            run.text = value
            value_run = run

        if suffix:
            suffix_run = _insert_run_after(paragraph, value_run, suffix)
            _apply_run_format(run, suffix_run)

        return value_run
    else:
        start_run = paragraph.runs[run_start_idx]
        end_run = paragraph.runs[run_end_idx]
        prefix = start_run.text[:run_start_offset]
        suffix = end_run.text[run_end_offset:]

        # Clear middle runs
        for i in range(run_start_idx + 1, run_end_idx):
            paragraph.runs[i].text = ''

        if prefix:
            start_run.text = prefix
            value_run = _insert_run_after(paragraph, start_run, value)
            _apply_run_format(start_run, value_run)
        else:
            start_run.text = value
            value_run = start_run

        if suffix:
            end_run.text = suffix
        else:
            end_run.text = ''

        return value_run


def _replace_placeholder_in_paragraph(paragraph, placeholder, value):
    """Find all occurrences of placeholder in paragraph and replace with value (run-level).
    Returns list of runs that were modified."""
    if placeholder not in paragraph.text:
        return []

    modified_runs = []
    full_text = ''.join(run.text for run in paragraph.runs)
    search_start = 0
    while True:
        idx = full_text.find(placeholder, search_start)
        if idx == -1:
            break
        end = idx + len(placeholder)
        result_run = _replace_span_in_paragraph(paragraph, idx, end, value)
        if result_run is None:
            break
        modified_runs.append(result_run)
        full_text = ''.join(run.text for run in paragraph.runs)
        search_start = idx + len(value)
    return modified_runs


def replace_placeholders(doc, data):
    """Replace placeholders in Shareholder Registry document using run-level replacement.
    Enforces 12pt Times New Roman on all replaced values."""
    print("===> Replacing placeholders in document...")

    company_name = data.get('companyName', '')
    formation_state = data.get('formationState', '')
    company_address = data.get('companyAddress', '')
    payment_date = data.get('paymentDate', '')
    authorized_shares = str(data.get('authorizedShares', '') or '')
    outstanding_shares = str(data.get('outstandingShares', '') or '')
    officer_1_name = data.get('officer1Name', '')
    officer_1_role = data.get('officer1Role', '')

    shareholders = data.get('shareholders', []) or []

    # Build placeholder map
    placeholders = {
        '{{Company Name}}': company_name,
        '{{Formation State}}': formation_state,
        '{{Company Address}}': company_address,
        '{{Payment Date}}': payment_date,
        '{{Authorized Shares}}': authorized_shares,
        '{{Outstanding Shares}}': outstanding_shares,
        '{{Officer 1 Name}}': officer_1_name,
        '{{Officer 1 Role}}': officer_1_role,
    }

    # Add shareholder placeholders (1-6)
    for idx in range(1, 7):
        num2 = f"{idx:02d}"
        shareholder = shareholders[idx - 1] if idx - 1 < len(shareholders) else {}
        placeholders[f'{{{{shareholder_{num2}_date}}}}'] = shareholder.get('date', '') or ''
        placeholders[f'{{{{shareholder_{num2}_name}}}}'] = shareholder.get('name', '') or ''
        placeholders[f'{{{{shareholder_{num2}_transaction}}}}'] = shareholder.get('transaction', '') or ''
        placeholders[f'{{{{shareholder_{num2}_shares}}}}'] = shareholder.get('shares', '') or ''
        placeholders[f'{{{{shareholder_{num2}_class}}}}'] = shareholder.get('class', '') or ''
        placeholders[f'{{{{shareholder_{num2}_percent}}}}'] = shareholder.get('percent', '') or ''

    def process_paragraph(paragraph):
        if '{{' not in paragraph.text:
            return
        modified_runs = []
        for ph, val in placeholders.items():
            modified_runs.extend(
                _replace_placeholder_in_paragraph(paragraph, ph, val)
            )
        # Enforce font on all modified runs
        for run in modified_runs:
            _enforce_font(run)

    # Process body paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)


def _set_table_col_widths(table, widths_inches):
    """Force column widths via direct XML: tblGrid, tblW, tblLayout, and per-cell tcW.
    widths_inches is a list of floats (inches per column)."""
    tbl = table._tbl

    # --- tblPr: fixed layout + total table width ---
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Fixed layout
    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)

    # Total table width (sum of columns) in twips (1 inch = 1440 twips)
    total_twips = sum(int(w * 1440) for w in widths_inches)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.insert(0, tblW)

    # --- tblGrid: define gridCol widths ---
    for old in tbl.findall(qn('w:tblGrid')):
        tbl.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_inches:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w * 1440)))
        tblGrid.append(gridCol)
    tblPr.addnext(tblGrid)

    # --- Per-cell tcW ---
    for row in table.rows:
        for i, w in enumerate(widths_inches):
            if i < len(row.cells):
                tc = row.cells[i]._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                for old in tcPr.findall(qn('w:tcW')):
                    tcPr.remove(old)
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(w * 1440)))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.insert(0, tcW)


def post_process_shareholder_registry(doc):
    """Best-practice formatting fixes for Shareholder Registry documents:
    1. Corporation Address: missing tab between label and value
    2. Remove "PAGE X" footer text
    3. Add vertical spacing above/below the shareholder table
    4. Adjust table column widths (wider Name column)
    """
    print("===> Post-processing: fixing template formatting...")

    # --- 1. Fix Corporation Address alignment ---
    # The header fields ("Corporation Address:", "Authorized Shares:", etc.) use
    # tab characters to align their values.  "Corporation Address:" is the longest
    # label, so a single default tab stop may leave its value misaligned with the
    # shorter labels.  Fix: ensure a tab exists between label and value AND set an
    # explicit tab stop at 2.5" so all header values line up at the same position.
    TAB_STOP_INCHES = 2.5  # position where header values should start
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if 'Corporation Address:' in text:
            runs = paragraph.runs
            # 1a. Insert a tab run between label and value if missing
            for i, run in enumerate(runs):
                if 'Corporation Address:' in run.text and i + 1 < len(runs):
                    next_run = runs[i + 1]
                    if not next_run.text.startswith('\t'):
                        tab_run = paragraph.add_run()
                        tab_el = OxmlElement('w:tab')
                        tab_run._element.append(tab_el)
                        run._element.addnext(tab_run._element)
                        print("===> Fixed: added tab between 'Corporation Address:' and value")
                    break

            # 1b. Set an explicit left tab stop so the value aligns with other fields
            pPr = paragraph._p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                paragraph._p.insert(0, pPr)
            # Remove any existing tabs element to avoid duplicates
            for old_tabs in pPr.findall(qn('w:tabs')):
                pPr.remove(old_tabs)
            tabs = OxmlElement('w:tabs')
            tab_stop = OxmlElement('w:tab')
            tab_stop.set(qn('w:val'), 'left')
            tab_stop.set(qn('w:pos'), str(int(TAB_STOP_INCHES * 1440)))  # twips
            tabs.append(tab_stop)
            pPr.append(tabs)
            print(f"===> Fixed: set tab stop at {TAB_STOP_INCHES}\" on Corporation Address paragraph")

        # Also normalise tab stops on the other header fields so everything is consistent
        if any(label in text for label in ['Authorized Shares:', 'Outstanding Shares:', 'Date of Formation:']):
            pPr = paragraph._p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                paragraph._p.insert(0, pPr)
            for old_tabs in pPr.findall(qn('w:tabs')):
                pPr.remove(old_tabs)
            tabs = OxmlElement('w:tabs')
            tab_stop = OxmlElement('w:tab')
            tab_stop.set(qn('w:val'), 'left')
            tab_stop.set(qn('w:pos'), str(int(TAB_STOP_INCHES * 1440)))
            tabs.append(tab_stop)
            pPr.append(tabs)
            print(f"===> Fixed: set tab stop at {TAB_STOP_INCHES}\" on '{text.split(':')[0].strip()}' paragraph")

    # --- 2. Remove "PAGE X" footer text ---
    pages_removed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r'^PAGE\s+\d+$', text):
            parent = paragraph._p.getparent()
            if parent is not None:
                parent.remove(paragraph._p)
                pages_removed += 1

    # --- 3. Add vertical spacing around the shareholder table ---
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == 'Common Stock':
            paragraph.paragraph_format.space_after = Pt(6)
            print("===> Fixed: added space_after on 'Common Stock' paragraph")
        if 'I hereby certify' in text:
            paragraph.paragraph_format.space_before = Pt(6)
            print("===> Fixed: added space_before on 'I hereby certify' paragraph")

    # --- 4. Adjust table column widths ---
    # Shareholder table has 6 columns:
    #   Date Acquired | Name | Transaction | # Shares | Class | Percentage
    # "Transaction" header must not wrap — needs ≥1.0" at 12pt TNR.
    # Give Name the most room; keep total ≈ 6.3" (letter page with ~1" margins).
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) == 6:
            col_widths_in = [
                0.95,   # Date Acquired (MM/DD/YYYY — needs ~0.95 to avoid wrap)
                1.60,   # Name (widest — room for full names)
                1.05,   # Transaction ("Transaction" header needs ~1.05 to avoid wrap)
                0.80,   # Number of Shares Owned
                0.85,   # Class of Shares
                1.05,   # Percentage Ownership
            ]  # total = 6.30"
            _set_table_col_widths(table, col_widths_in)
            print("===> Fixed: adjusted shareholder table column widths")

    print(f"===> Post-processing done: {pages_removed} PAGE X removed")


def lambda_handler(event, context):
    print("===> Shareholder Registry Lambda invoked")

    try:
        body = json.loads(event.get('body', '{}'))
    except Exception:
        body = event if isinstance(event, dict) else {}

    form_data = body.get('form_data', {}) or {}
    s3_bucket = body.get('s3_bucket', OUTPUT_BUCKET)
    s3_key = body.get('s3_key', '')
    template_url = body.get('templateUrl')
    return_docx = body.get('return_docx', False)

    if not template_url:
        return {
            "statusCode": 400,
            "body": json.dumps({"error": "Missing 'templateUrl'"})
        }

    template_bucket, template_key = extract_s3_info(template_url)
    if not template_bucket or not template_key:
        template_bucket = BUCKET_NAME
        template_key = 'templates/shareholder-registry/shareholder-registry-1.docx'
        print(f"===> Could not parse templateUrl, using default: s3://{template_bucket}/{template_key}")

    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "template_shareholder_registry.docx")
        output_path = os.path.join(tmpdir, "filled_shareholder_registry.docx")

        download_from_s3(template_bucket, template_key, template_path)

        doc = Document(template_path)
        replace_placeholders(doc, form_data)
        post_process_shareholder_registry(doc)
        doc.save(output_path)

        upload_to_s3(output_path, s3_bucket, s3_key)

        if return_docx:
            with open(output_path, "rb") as f:
                docx_content = f.read()
            encoded = base64.b64encode(docx_content).decode('utf-8')
            return {
                "statusCode": 200,
                "headers": {"Content-Type": "application/json"},
                "body": json.dumps({
                    "message": "Shareholder Registry generated successfully",
                    "docx_base64": encoded
                })
            }

        return {
            "statusCode": 200,
            "body": json.dumps({"message": "Shareholder Registry generated successfully"})
        }

    return {
        "statusCode": 500,
        "body": json.dumps({"error": "Failed to generate Shareholder Registry"})
    }
