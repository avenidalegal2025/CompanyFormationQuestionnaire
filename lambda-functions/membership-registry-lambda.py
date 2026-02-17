import os
import json
import tempfile
import boto3
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import base64

# Constants
# Template bucket (where templates are stored)
TEMPLATE_BUCKET = os.environ.get('TEMPLATE_BUCKET', 'company-formation-template-llc-and-inc')
# Output bucket (where filled documents are saved)
OUTPUT_BUCKET = os.environ.get('OUTPUT_BUCKET', 'avenida-legal-documents')
# Legacy support - use template bucket as fallback
BUCKET_NAME = os.environ.get('BUCKET_NAME', TEMPLATE_BUCKET)

# Initialize S3 client
s3_client = boto3.client('s3', region_name='us-west-1')

def download_from_s3(bucket, key, local_path):
    """Download file from S3 to local path"""
    print(f"===> Downloading s3://{bucket}/{key} to {local_path}")
    s3_client.download_file(bucket, key, local_path)
    print(f"===> Downloaded successfully")

def upload_to_s3(local_path, bucket, key):
    """Upload file from local path to S3"""
    print(f"===> Uploading {local_path} to s3://{bucket}/{key}")
    s3_client.upload_file(local_path, bucket, key)
    print(f"===> Uploaded successfully")

def extract_s3_info(url):
    """
    Extract bucket and key from S3 URL
    Supports formats:
    - https://bucket.s3.region.amazonaws.com/key
    - s3://bucket/key
    """
    if not url:
        return None, None
    
    # Handle s3:// URLs
    if url.startswith('s3://'):
        parts = url[5:].split('/', 1)
        if len(parts) == 2:
            return parts[0], parts[1]
        return parts[0], ''
    
    # Handle https:// URLs
    if url.startswith('https://'):
        # Pattern: https://bucket.s3.region.amazonaws.com/key
        match = re.match(r'https://([^.]+)\.s3[^/]*\.amazonaws\.com/(.+)', url)
        if match:
            return match.group(1), match.group(2)
        
        # Pattern: https://s3.region.amazonaws.com/bucket/key
        match = re.match(r'https://s3[^/]*\.amazonaws\.com/([^/]+)/(.+)', url)
        if match:
            return match.group(1), match.group(2)
    
    return None, None

def format_percentage(value):
    """Format ownership percentage - pixel perfect formatting (no trailing zeros)"""
    if value is None:
        return "0%"
    
    # Convert to float
    if isinstance(value, (int, float)):
        num = float(value)
    elif isinstance(value, str):
        try:
            # Handle comma as decimal separator (e.g., "50,1" -> 50.1)
            num = float(value.replace(',', '.'))
        except:
            return "0%"
    else:
        return "0%"
    
    # If value is between 0 and 1, multiply by 100 (Airtable stores as decimal)
    if 0 <= num <= 1:
        num = num * 100
    
    # Format without trailing zeros - pixel perfect
    # Remove trailing zeros and decimal point if not needed
    if num == int(num):
        # Whole number: "50%" not "50.00%"
        return f"{int(num)}%"
    else:
        # Has decimals: remove trailing zeros
        # "50.1%" not "50.10%", "50.25%" not "50.250%"
        formatted = f"{num:.10f}".rstrip('0').rstrip('.')
        return f"{formatted}%"

def format_address(address_str):
    """Format address string, handling None and empty values"""
    if not address_str or address_str.strip() == '':
        return ''
    return address_str.strip()


def _formation_date_to_mm_dd_yyyy(value: str) -> str:
    """Convert formation date (long or raw) to MM/DD/YYYY for Particulars of ownership / Date Acquired."""
    if not value or not value.strip():
        return ''
    # Already MM/DD/YYYY or M/D/YYYY
    m = re.match(r'^\s*(\d{1,2})/(\d{1,2})/(\d{4})\s*$', value)
    if m:
        return f"{int(m.group(1)):02d}/{int(m.group(2)):02d}/{m.group(3)}"
    # ISO
    m = re.match(r'^\s*(\d{4})-(\d{1,2})-(\d{1,2})\s*', value)
    if m:
        return f"{int(m.group(2)):02d}/{int(m.group(3)):02d}/{m.group(1)}"
    # "8th day of February, 2026" or "February 8th, 2026"
    months = {'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
              'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12}
    m = re.search(r'(\d{1,2})(?:st|nd|rd|th)?\s+day of\s+(\w+)\s*,\s*(\d{4})', value, re.IGNORECASE)
    if m:
        day, month_name, year = int(m.group(1)), months.get(m.group(2).lower()), m.group(3)
        if month_name:
            return f"{month_name:02d}/{day:02d}/{year}"
    m = re.search(r'(\w+)\s+(\d{1,2})(?:st|nd|rd|th)?\s*,\s*(\d{4})', value, re.IGNORECASE)
    if m:
        month_name, day, year = months.get(m.group(1).lower()), int(m.group(2)), m.group(3)
        if month_name:
            return f"{month_name:02d}/{day:02d}/{year}"
    return value

def _set_cell_text_preserving_format(cell, new_text):
    """
    Set cell text while preserving the first paragraph's run formatting.
    Falls back to cell.text if no runs exist.
    """
    if cell.paragraphs and cell.paragraphs[0].runs:
        first_para = cell.paragraphs[0]
        first_para.runs[0].text = new_text
        for run in first_para.runs[1:]:
            run.text = ''
        # Clear any extra paragraphs
        for para in cell.paragraphs[1:]:
            for run in para.runs:
                run.text = ''
    else:
        cell.text = new_text


def _replace_in_paragraph_preserving_format(paragraph, replace_fn):
    """
    Replace placeholders inside a paragraph while preserving run-level formatting
    (font, size, bold, italic, color, etc.).

    Word splits text across multiple runs, so a placeholder like {{COMPANY_NAME}}
    might be split as: run1="{{COMPANY", run2="_NAME}}".

    Strategy:
    1. Join all run texts, apply replacements on the joined string.
    2. If the joined result differs, redistribute the new text across the original runs
       keeping each run's formatting intact.
    """
    if not paragraph.runs:
        # No runs — fall back to direct text replacement (rare, but safe)
        original = paragraph.text
        new_text = replace_fn(original)
        if new_text != original:
            paragraph.text = new_text
        return

    # Concatenate all run texts
    full_original = ''.join(run.text for run in paragraph.runs)
    full_replaced = replace_fn(full_original)

    if full_replaced == full_original:
        return  # Nothing changed

    # Put all new text in the first run, clear the rest — this preserves
    # the first run's formatting (font, size, bold, etc.) for the whole paragraph.
    # This is safe because placeholders typically share the same formatting as
    # surrounding text in legal document templates.
    paragraph.runs[0].text = full_replaced
    for run in paragraph.runs[1:]:
        run.text = ''


def replace_placeholders(doc, data):
    """
    Replace placeholders in Word document with actual data.
    Preserves run-level formatting (font, size, bold, italic) to prevent
    font corruption when converting to PDF via LibreOffice.

    Supports two styles of templates:
      1) Generic placeholders: {{COMPANY_NAME}}, {{COMPANY_ADDRESS}}, {{FORMATION_STATE}}, {{FORMATION_DATE}}
      2) Legacy membership templates with placeholders like:
         {{llc_name_text}}, {{full_llc_address}}, {{full_state}}, {{full_state_caps}},
         {{Date_of_formation_LLC}}, {{member_01_full_name}}, {{member_01_pct}}, {{manager_01_full_name}}, etc.
    """
    print("===> Replacing placeholders in document...")

    # Company information
    company_name = data.get('companyName', '')
    company_address = format_address(data.get('companyAddress', ''))
    formation_state = data.get('formationState', '')
    formation_date = data.get('formationDate', '')
    # Numeric date (MM/DD/YYYY) for Particulars of ownership / Date Acquired only
    formation_date_numeric = data.get('formationDateNumeric') or _formation_date_to_mm_dd_yyyy(formation_date)

    # Members / managers
    members = data.get('members', []) or []
    managers = data.get('managers', []) or []
    print(f"===> Found {len(members)} members and {len(managers)} managers in form data")

    # Helper to replace all known placeholders in a text string.
    # In table cells (Particulars of ownership) use numeric date (MM/DD/YYYY) for date placeholders.
    # When as_of_the_date is True (e.g. certification "as of ... date"), use "the 9th day..." not "9th day...".
    def replace_in_text(text: str, in_table_cell: bool = False, as_of_the_date: bool = False) -> str:
        if not text:
            return text
        if in_table_cell:
            date_val = formation_date_numeric
        else:
            # Certification sentence: "as of the 9th day of February, 2026"
            date_val = ('the ' + formation_date) if (as_of_the_date and formation_date) else formation_date
        # Generic placeholders
        text = text.replace('{{COMPANY_NAME}}', company_name)
        text = text.replace('{{COMPANY_ADDRESS}}', company_address)
        text = text.replace('{{FORMATION_STATE}}', formation_state)
        text = text.replace('{{FORMATION_DATE}}', date_val)
        # Legacy LLC-specific placeholders
        text = text.replace('{{llc_name_text}}', company_name)
        text = text.replace('{{full_llc_address}}', company_address)
        text = text.replace('{{full_state}}', formation_state)
        text = text.replace('{{full_state_caps}}', formation_state.upper() if formation_state else '')
        text = text.replace('{{Date_of_formation_LLC}}', date_val)

        # Member placeholders: {{member_01_full_name}}, {{member_01_pct}}, etc.
        for idx, member in enumerate(members, start=1):
            num2 = f"{idx:02d}"
            # Names
            for ph in (f'{{{{member_{num2}_full_name}}}}', f'{{{{member_{idx}_full_name}}}}'):
                if ph in text:
                    text = text.replace(ph, member.get('name', ''))
            # Ownership percent
            pct_val = member.get('ownershipPercent', 0) or 0
            try:
                pct_num = float(pct_val)
                if 0 < pct_num <= 1:
                    pct_num = pct_num * 100.0
            except Exception:
                pct_num = 0
            pct_str = f"{pct_num:.2f}".rstrip('0').rstrip('.') if pct_num else "0"
            for ph in (f'{{{{member_{num2}_pct}}}}', f'{{{{member_{idx}_pct}}}}'):
                if ph in text:
                    text = text.replace(ph, pct_str)

        # Manager placeholders: {{manager_01_full_name}}, etc.
        for idx, manager in enumerate(managers, start=1):
            num2 = f"{idx:02d}"
            for ph in (f'{{{{manager_{num2}_full_name}}}}', f'{{{{manager_{idx}_full_name}}}}'):
                if ph in text:
                    text = text.replace(ph, manager.get('name', ''))

        return text

    # First pass: replace in paragraphs (preserving formatting)
    for paragraph in doc.paragraphs:
        original = paragraph.text
        has_as_of = 'as of' in original.lower()
        has_date_placeholder = '{{FORMATION_DATE}}' in original or '{{Date_of_formation_LLC}}' in original
        as_of_the_date = has_as_of and has_date_placeholder
        _aod = as_of_the_date  # capture for closure
        _replace_in_paragraph_preserving_format(
            paragraph,
            lambda text, _a=_aod: replace_in_text(text, in_table_cell=False, as_of_the_date=_a)
        )

    # Second pass: replace in tables (cell paragraphs) — use numeric date in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph_preserving_format(
                        paragraph,
                        lambda text: replace_in_text(text, in_table_cell=True)
                    )
    
    # Existing table-based logic for templates that use dynamic rows
    print("===> Applying table-based member/manager filling logic (if applicable)")
    print(f"===> Searching tables for header-based member/manager layouts")
    
    # Find tables with member and manager information
    # Typically, membership registry has tables with columns: Name, Address, Ownership %
    for table in doc.tables:
        # Check if this is the members table (usually has headers like "Name", "Address", "Ownership")
        is_member_table = False
        if len(table.rows) > 0:
            header_row = table.rows[0]
            header_text = ' '.join([cell.text.strip().lower() for cell in header_row.cells])
            if 'name' in header_text and ('address' in header_text or 'ownership' in header_text):
                is_member_table = True
        
        if is_member_table:
            print(f"===> Found member table with {len(table.rows)} rows")
            
            # Map column headers to indices
            header_row = table.rows[0]
            column_map = {}
            for idx, cell in enumerate(header_row.cells):
                header_text = cell.text.strip().lower()
                if 'member' in header_text and 'name' in header_text:
                    column_map['name'] = idx
                elif 'date' in header_text and 'acquired' in header_text:
                    column_map['date'] = idx
                elif 'address' in header_text:
                    column_map['address'] = idx
                elif 'ownership' in header_text and 'percent' in header_text:
                    column_map['ownership'] = idx
                elif 'percentage' in header_text and 'ownership' in header_text:
                    column_map['ownership'] = idx
                elif 'transaction' in header_text:
                    column_map['transaction'] = idx
                elif 'ssn' in header_text or 'social' in header_text:
                    column_map['ssn'] = idx
            
            print(f"===> Column mapping: {column_map}")
            
            # Add member rows
            for i, member in enumerate(members):
                # Add a new row if needed
                if i + 1 >= len(table.rows):
                    table.add_row()
                
                row = table.rows[i + 1]  # Skip header row (index 0)
                
                # Fill member data based on column mapping (preserve formatting)
                if 'name' in column_map and len(row.cells) > column_map['name']:
                    _set_cell_text_preserving_format(row.cells[column_map['name']], member.get('name', ''))

                if 'address' in column_map and len(row.cells) > column_map['address']:
                    _set_cell_text_preserving_format(row.cells[column_map['address']], format_address(member.get('address', '')))

                if 'ownership' in column_map and len(row.cells) > column_map['ownership']:
                    _set_cell_text_preserving_format(row.cells[column_map['ownership']], format_percentage(member.get('ownershipPercent', 0)))

                if 'transaction' in column_map and len(row.cells) > column_map['transaction']:
                    if 'ownership' not in column_map:
                        _set_cell_text_preserving_format(row.cells[column_map['transaction']], format_percentage(member.get('ownershipPercent', 0)))

                if 'date' in column_map and len(row.cells) > column_map['date']:
                    _set_cell_text_preserving_format(row.cells[column_map['date']], formation_date_numeric)

                if 'ssn' in column_map and len(row.cells) > column_map['ssn']:
                    ssn = member.get('ssn', '')
                    if ssn and ssn.upper() not in ['N/A', 'N/A-FOREIGN', '']:
                        _set_cell_text_preserving_format(row.cells[column_map['ssn']], ssn)
                    else:
                        _set_cell_text_preserving_format(row.cells[column_map['ssn']], 'N/A')
            
            # Remove extra rows if we have fewer members than rows
            while len(table.rows) > len(members) + 1:  # +1 for header
                table._element.remove(table.rows[-1]._element)
            
            # Now handle manager table (if exists)
            # Look for a table with "Manager" in header
            for table in doc.tables:
                if len(table.rows) > 0:
                    header_row = table.rows[0]
                    header_text = ' '.join([cell.text.strip().lower() for cell in header_row.cells])
                    if 'manager' in header_text and 'name' in header_text:
                        print(f"===> Found manager table with {len(table.rows)} rows")
                        # Add manager rows
                        for i, manager in enumerate(managers):
                            # Add a new row if needed
                            if i + 1 >= len(table.rows):
                                table.add_row()
                            
                            row = table.rows[i + 1]  # Skip header row (index 0)
                            
                            # Fill manager data (preserve formatting)
                            # Column 0: Manager Name
                            if len(row.cells) > 0:
                                _set_cell_text_preserving_format(row.cells[0], manager.get('name', ''))

                            # Column 1: Address
                            if len(row.cells) > 1:
                                _set_cell_text_preserving_format(row.cells[1], format_address(manager.get('address', '')))
                        
                        # Remove extra rows if we have fewer managers than rows
                        while len(table.rows) > len(managers) + 1:  # +1 for header
                            table._element.remove(table.rows[-1]._element)
                        break
            
            break
    
    # Final pass: any remaining placeholders in table cells (preserving formatting)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = paragraph.text
                    if '{{COMPANY_NAME}}' in full_text or '{{COMPANY_ADDRESS}}' in full_text or '{{FORMATION_STATE}}' in full_text or '{{FORMATION_DATE}}' in full_text or '{{Date_of_formation_LLC}}' in full_text:
                        def _final_replace(text):
                            text = text.replace('{{COMPANY_NAME}}', company_name)
                            text = text.replace('{{COMPANY_ADDRESS}}', company_address)
                            text = text.replace('{{FORMATION_STATE}}', formation_state)
                            text = text.replace('{{FORMATION_DATE}}', formation_date_numeric)
                            text = text.replace('{{Date_of_formation_LLC}}', formation_date_numeric)
                            return text
                        _replace_in_paragraph_preserving_format(paragraph, _final_replace)
    
    print("===> Placeholders replaced successfully")

def lambda_handler(event, context):
    print("===> RAW EVENT:")
    print(json.dumps(event))
    
    try:
        # Parse request body
        if "body" in event:
            body = json.loads(event["body"])
        else:
            body = event
        
        form_data = body.get("form_data")
        s3_bucket = body.get("s3_bucket")
        s3_key = body.get("s3_key")
        templateUrl = body.get("templateUrl")
        return_docx = body.get("return_docx", False)
        
        if not form_data:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Missing 'form_data'"})
            }
        
        if not s3_bucket or not s3_key:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Missing 's3_bucket' or 's3_key'"})
            }
        
        if not templateUrl:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Missing 'templateUrl'"})
            }
        
        # Extract template bucket and key from URL
        template_bucket, template_key = extract_s3_info(templateUrl)
        if not template_bucket or not template_key:
            # Fallback to template bucket with default path
            template_bucket = TEMPLATE_BUCKET
            template_key = 'llc-formation-templates/membership-registry-all-templates/membership-registry-1-member/Template Membership Registry_1 Members_1 Manager.docx'
            print(f"===> Could not parse templateUrl, using default: s3://{template_bucket}/{template_key}")
        else:
            print(f"===> Template: s3://{template_bucket}/{template_key}")
        
        print(f"===> Output: s3://{s3_bucket}/{s3_key}")
        
    except Exception as e:
        return {
            "statusCode": 400,
            "body": json.dumps({"error": "Invalid input payload", "details": str(e)})
        }
    
    # Prepare file paths
    tmpdir = tempfile.gettempdir()
    template_path = os.path.join(tmpdir, "template_membership_registry.docx")
    output_path = os.path.join(tmpdir, "filled_membership_registry.docx")
    
    try:
        # Download template from S3
        download_from_s3(template_bucket, template_key, template_path)
        
        # Load Word document
        print("===> Loading Word document...")
        doc = Document(template_path)
        
        # Replace placeholders with form data
        replace_placeholders(doc, form_data)
        
        # Save filled document
        print("===> Saving filled document...")
        doc.save(output_path)
        print(f"===> Saved to {output_path}")
        
        # Upload to S3
        upload_to_s3(output_path, s3_bucket, s3_key)
        
        # Return response
        if return_docx:
            # Read file and return as base64-encoded binary (required by Lambda proxy)
            with open(output_path, 'rb') as f:
                docx_content = f.read()
            
            encoded = base64.b64encode(docx_content).decode('utf-8')
            
            return {
                "statusCode": 200,
                "headers": {
                    "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "Content-Disposition": f'attachment; filename="{os.path.basename(s3_key)}"'
                },
                "body": encoded,
                "isBase64Encoded": True
            }
        else:
            return {
                "statusCode": 200,
                "body": json.dumps({
                    "message": "✅ Document uploaded to S3",
                    "s3_bucket": s3_bucket,
                    "s3_key": s3_key,
                    "s3_url": f"s3://{s3_bucket}/{s3_key}"
                })
            }
    
    except Exception as e:
        import traceback
        error_msg = str(e)
        traceback.print_exc()
        print(f"===> ERROR: {error_msg}")
        return {
            "statusCode": 500,
            "body": json.dumps({
                "error": "Failed to generate membership registry",
                "details": error_msg
            })
        }
