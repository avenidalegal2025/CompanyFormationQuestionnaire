import os
import json
import tempfile
import boto3
import re
import base64
from copy import deepcopy
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Constants
TEMPLATE_BUCKET = os.environ.get('TEMPLATE_BUCKET', 'avenida-legal-documents')
OUTPUT_BUCKET = os.environ.get('OUTPUT_BUCKET', 'avenida-legal-documents')
BUCKET_NAME = os.environ.get('BUCKET_NAME', TEMPLATE_BUCKET)

# Initialize S3 client
s3_client = boto3.client('s3', region_name=os.environ.get('AWS_REGION', 'us-west-1'))


# =============================================================================
#  S3 helpers (identical to other lambdas)
# =============================================================================

def download_from_s3(bucket, key, local_path):
    """Download file from S3 to local path"""
    print(f"===> Downloading s3://{bucket}/{key} to {local_path}")
    s3_client.download_file(bucket, key, local_path)
    print("===> Downloaded successfully")


def upload_to_s3(local_path, bucket, key):
    """Upload file from local path to S3"""
    print(f"===> Uploading {local_path} to s3://{bucket}/{key}")
    s3_client.upload_file(local_path, bucket, key)
    print("===> Uploaded successfully")


def extract_s3_info(url):
    """Extract bucket and key from S3 URL"""
    if not url:
        return None, None

    if url.startswith('s3://'):
        parts = url[5:].split('/', 1)
        if len(parts) == 2:
            return parts[0], parts[1]
        return parts[0], ''

    if url.startswith('https://'):
        match = re.match(r'https://([^.]+)\.s3[^/]*\.amazonaws\.com/(.+)', url)
        if match:
            return match.group(1), match.group(2)
        match = re.match(r'https://s3[^/]*\.amazonaws\.com/([^/]+)/(.+)', url)
        if match:
            return match.group(1), match.group(2)

    return None, None


# =============================================================================
#  Formatting helpers
# =============================================================================

def format_percentage(value):
    """Format ownership percentage — no trailing zeros."""
    if value is None:
        return "0%"
    if isinstance(value, (int, float)):
        num = float(value)
    elif isinstance(value, str):
        try:
            num = float(value.replace(',', '.'))
        except Exception:
            return "0%"
    else:
        return "0%"
    if 0 < num <= 1:
        num = num * 100
    if num == int(num):
        return f"{int(num)}%"
    else:
        formatted = f"{num:.10f}".rstrip('0').rstrip('.')
        return f"{formatted}%"


def format_currency(value):
    """Format a number as $X,XXX.00"""
    if value is None:
        return "$0.00"
    try:
        num = float(value)
    except Exception:
        return "$0.00"
    return f"${num:,.2f}"


def format_shares(value):
    """Format share count with thousands separator."""
    if value is None:
        return "0"
    try:
        return f"{int(value):,}"
    except Exception:
        return str(value)


def _parse_date(value):
    """Parse m/d/yyyy or yyyy-mm-dd; return (day, month_name, year) or None."""
    if not value or not value.strip():
        return None
    match = (re.match(r'^\s*(\d{1,2})/(\d{1,2})/(\d{4})\s*$', value) or
             re.match(r'^\s*(\d{4})-(\d{1,2})-(\d{1,2})\s*$', value))
    if not match:
        return None
    if len(match.group(1)) == 4:
        year, month, day = int(match.group(1)), int(match.group(2)), int(match.group(3))
    else:
        month, day, year = int(match.group(1)), int(match.group(2)), int(match.group(3))
    try:
        date = datetime(year, month, day)
        return (day, date.strftime('%B'), year)
    except ValueError:
        return None


def _ordinal_suffix(day):
    if 11 <= day <= 13:
        return 'th'
    return {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')


def _number_to_word(n):
    words = {
        1: 'first', 2: 'second', 3: 'third', 4: 'fourth', 5: 'fifth', 6: 'sixth', 7: 'seventh',
        8: 'eighth', 9: 'ninth', 10: 'tenth', 11: 'eleventh', 12: 'twelfth', 13: 'thirteenth',
        14: 'fourteenth', 15: 'fifteenth', 16: 'sixteenth', 17: 'seventeenth', 18: 'eighteenth',
        19: 'nineteenth', 20: 'twentieth', 21: 'twenty-first', 22: 'twenty-second', 23: 'twenty-third',
        24: 'twenty-fourth', 25: 'twenty-fifth', 26: 'twenty-sixth', 27: 'twenty-seventh',
        28: 'twenty-eighth', 29: 'twenty-ninth', 30: 'thirtieth', 31: 'thirty-first',
    }
    return words.get(n, str(n))


def format_legal_date(value):
    """Body text date: 'eighth day of February, 2026'."""
    if not value:
        return value
    if 'day of' in value:
        return value.strip()
    parsed = _parse_date(value)
    if not parsed:
        return value
    day, month_name, year = parsed
    return f"{_number_to_word(day)} day of {month_name}, {year}"


def format_witness_date(value):
    """IN WITNESS WHEREOF date: '8th day of February, 2026'."""
    if not value:
        return value
    if 'day of' in value:
        return value.strip()
    parsed = _parse_date(value)
    if not parsed:
        return value
    day, month_name, year = parsed
    return f"{day}{_ordinal_suffix(day)} day of {month_name}, {year}"


# =============================================================================
#  Run-level replacement engine (from org-resolution lambda)
# =============================================================================

def _insert_run_after(paragraph, run, text):
    """Insert a new run after the given run, returning the new run."""
    new_run = paragraph.add_run(text)
    run._element.addnext(new_run._element)
    return new_run


def _apply_run_format(source_run, target_run, bold_override=None):
    """Copy formatting from source_run to target_run."""
    if source_run._element.rPr is not None:
        target_run._element.insert(0, deepcopy(source_run._element.rPr))
    if bold_override is not None:
        target_run.bold = bold_override


def _enforce_font(run):
    """Ensure a modified run uses 12pt Times New Roman."""
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def _replace_span_in_paragraph(paragraph, start, end, value, bold_value=False):
    """Replace a character span [start, end) in the paragraph's runs with value."""
    pos = 0
    run_start_idx = None
    run_end_idx = None
    run_start_offset = 0
    run_end_offset = 0

    for i, run in enumerate(paragraph.runs):
        run_len = len(run.text)
        if run_start_idx is None and pos + run_len > start:
            run_start_idx = i
            run_start_offset = start - pos
        if run_start_idx is not None and pos + run_len >= end:
            run_end_idx = i
            run_end_offset = end - pos
            break
        pos += run_len

    if run_start_idx is None or run_end_idx is None:
        return None

    if not bold_value:
        # Simple replacement — no bold split needed
        if run_start_idx == run_end_idx:
            run = paragraph.runs[run_start_idx]
            run.text = run.text[:run_start_offset] + value + run.text[run_end_offset:]
            return run
        else:
            start_run = paragraph.runs[run_start_idx]
            end_run = paragraph.runs[run_end_idx]
            start_run.text = start_run.text[:run_start_offset] + value
            for i in range(run_start_idx + 1, run_end_idx):
                paragraph.runs[i].text = ''
            end_run.text = end_run.text[run_end_offset:]
            return start_run
    else:
        # Bold the replacement value
        if run_start_idx == run_end_idx:
            run = paragraph.runs[run_start_idx]
            prefix = run.text[:run_start_offset]
            suffix = run.text[run_end_offset:]

            if prefix:
                run.text = prefix
                run.bold = False
                value_run = _insert_run_after(paragraph, run, value)
                _apply_run_format(run, value_run, True)
            else:
                run.text = value
                run.bold = True
                value_run = run

            if suffix:
                suffix_run = _insert_run_after(paragraph, value_run, suffix)
                _apply_run_format(run, suffix_run, False)

            return value_run
        else:
            start_run = paragraph.runs[run_start_idx]
            end_run = paragraph.runs[run_end_idx]
            prefix = start_run.text[:run_start_offset]
            suffix = end_run.text[run_end_offset:]

            for i in range(run_start_idx + 1, run_end_idx):
                paragraph.runs[i].text = ''

            if prefix:
                start_run.text = prefix
                start_run.bold = False
                value_run = _insert_run_after(paragraph, start_run, value)
                _apply_run_format(start_run, value_run, True)
            else:
                start_run.text = value
                start_run.bold = True
                value_run = start_run

            if suffix:
                end_run.text = suffix
                end_run.bold = False
            else:
                end_run.text = ''

            return value_run


def _replace_placeholder_in_paragraph(paragraph, placeholder, value, bold_value=False):
    """Find all occurrences of placeholder in paragraph and replace with value (run-level)."""
    if placeholder not in paragraph.text:
        return []
    modified_runs = []
    full_text = ''.join(run.text for run in paragraph.runs)
    search_start = 0
    while True:
        idx = full_text.find(placeholder, search_start)
        if idx == -1:
            break
        end = idx + len(placeholder)
        result_run = _replace_span_in_paragraph(paragraph, idx, end, value, bold_value)
        if result_run is None:
            break
        modified_runs.append(result_run)
        full_text = ''.join(run.text for run in paragraph.runs)
        search_start = idx + len(value)
    return modified_runs


def _replace_pct_placeholder_in_paragraph(paragraph, placeholder, pct_with_percent, pct_without_percent):
    """Smart percentage replacement: if template already has % after placeholder, use value without %."""
    if placeholder not in paragraph.text:
        return
    full_text = ''.join(run.text for run in paragraph.runs)
    search_start = 0
    while True:
        idx = full_text.find(placeholder, search_start)
        if idx == -1:
            break
        end = idx + len(placeholder)
        has_percent = end < len(full_text) and full_text[end] == '%'
        value = pct_without_percent if has_percent else pct_with_percent
        result = _replace_span_in_paragraph(paragraph, idx, end, value)
        if result is None:
            break
        full_text = ''.join(run.text for run in paragraph.runs)
        search_start = idx + len(value)


# =============================================================================
#  Voting text logic
# =============================================================================

# Each tuple: (form_data key, unique phrase in template, which word to replace)
# The phrases below are unique in the Shareholder Agreement template so we can
# safely do a single-occurrence targeted replacement without section markers.

VOTING_PHRASES = [
    # Sec 3.2(a) — dissolution
    ("dissolutionVoting",
     "Majority election to dissolve by the Shareholders",
     "Majority"),
    # Sec 4.3 — new shareholders
    ("newShareholderVoting",
     "approved by a Majority of the Shareholders, and following the restrictions",
     "Majority"),
    # Sec 4.5 — additional capital contributions
    ("capitalCallsVoting",
     "with the Majority approval of the Board of Directors. No Shareholder",
     "Majority"),
    # Sec 7.3 — shareholder loans
    ("shareholderLoansVoting",
     "without the explicit Majority approval of the Board of Directors",
     "Majority"),
    # Sec 9.1(v) — sale of entire company
    ("saleOfCompanyVoting",
     "unless the Majority consent or approval of both the Shareholders and the Board of Directors is given",
     "Majority"),
    # Sec 10.1 — major decisions / limitation on officers
    ("majorDecisionsVoting",
     "Without the Majority consent of the Board of Directors, the Officers shall not have authority to",
     "Majority"),
    # Sec 12.1 — removal of officers/directors
    ("removalVoting",
     "by the Majority vote of the Shareholders at a meeting called expressly for that purpose",
     "Majority"),
    # Sec 9.1(iv) — divorce transfer consent
    ("divorceTransferVoting",
     "Majority of the Shareholders shall consent to such transfer",
     "Majority"),
]


def _get_voting_label(value, super_majority_pct=None):
    """Convert form value to display text for the agreement.
    'majority' → 'Majority'
    'supermajority' → 'Super Majority'
    'unanimous' → 'Unanimous'
    """
    if not value:
        return 'Majority'
    v = value.strip().lower()
    if v == 'unanimous':
        return 'Unanimous'
    if v in ('supermajority', 'super_majority', 'super majority'):
        return 'Super Majority'
    return 'Majority'


def apply_voting_replacements(doc, data):
    """Apply section-specific voting text replacements.
    Each voting field maps to a unique phrase in the template.
    We replace 'Majority' within that phrase with the configured value.
    """
    print("===> Applying voting text replacements...")
    super_majority_pct = data.get('superMajorityThreshold')
    replacements_made = 0

    for field_key, target_phrase, old_word in VOTING_PHRASES:
        voting_value = data.get(field_key, 'majority')
        new_word = _get_voting_label(voting_value, super_majority_pct)

        if new_word == old_word:
            continue  # No change needed

        new_phrase = target_phrase.replace(old_word, new_word, 1)

        for paragraph in doc.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            if target_phrase in full_text:
                idx = full_text.find(target_phrase)
                end = idx + len(target_phrase)
                result = _replace_span_in_paragraph(paragraph, idx, end, new_phrase)
                if result:
                    _enforce_font(result)
                    replacements_made += 1
                    print(f"===> Voting: {field_key} → '{new_word}' in: {target_phrase[:60]}...")
                break

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = ''.join(run.text for run in paragraph.runs)
                        if target_phrase in full_text:
                            idx = full_text.find(target_phrase)
                            end = idx + len(target_phrase)
                            result = _replace_span_in_paragraph(paragraph, idx, end, new_phrase)
                            if result:
                                _enforce_font(result)
                                replacements_made += 1

    print(f"===> Voting replacements done: {replacements_made} changes")


# =============================================================================
#  Majority / Super Majority definition replacement (Sec 1.6 / 1.11)
# =============================================================================

def apply_majority_definition(doc, data):
    """Update the Majority definition in Sec 1.6 if threshold differs from 50%.
    Add Super Majority definition at Sec 1.11 if supermajority is used."""
    majority_threshold = data.get('majorityThreshold', 50)
    super_majority_pct = data.get('superMajorityThreshold')

    # Sec 1.6: "FIFTY PERCENT (50.00%)" — replace if threshold changed
    if majority_threshold and majority_threshold != 50:
        pct_str = f"{majority_threshold:.2f}".rstrip('0').rstrip('.')
        # Build the word form
        # For simplicity, use numeric form for non-standard thresholds
        old_text = "FIFTY PERCENT (50.00%)"
        new_text = f"{pct_str} PERCENT ({pct_str}%)"
        for paragraph in doc.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            if old_text in full_text:
                idx = full_text.find(old_text)
                _replace_span_in_paragraph(paragraph, idx, idx + len(old_text), new_text)
                print(f"===> Updated Majority definition to {pct_str}%")
                break

    # Sec 1.11: Super Majority definition — the template has a blank:
    # 'greater than ___%'
    # Fill in the super majority percentage
    if super_majority_pct:
        pct_str = f"{super_majority_pct:.2f}".rstrip('0').rstrip('.')
        for paragraph in doc.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            if 'greater than __' in full_text or 'greater than ___' in full_text:
                # Replace the blank with the percentage
                old = re.search(r'greater than _+%?', full_text)
                if old:
                    new_text = f"greater than {pct_str}%"
                    _replace_span_in_paragraph(paragraph, old.start(), old.end(), new_text)
                    print(f"===> Filled Super Majority definition: {pct_str}%")
                break


# =============================================================================
#  Bank account signature replacement (Sec 10.7)
# =============================================================================

def apply_bank_signature_replacement(doc, data):
    """Replace bank signature requirement in Sec 10.7.
    Default: 'upon the signature of one of the Officers'
    If bankSignatures == 2: 'upon the signature of two of the Officers'
    """
    bank_sigs = data.get('bankSignatures', 1)
    if bank_sigs != 2:
        return

    old_text = "upon the signature of one of the Officers"
    new_text = "upon the signature of two of the Officers"

    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        if old_text in full_text:
            idx = full_text.find(old_text)
            result = _replace_span_in_paragraph(paragraph, idx, idx + len(old_text), new_text)
            if result:
                _enforce_font(result)
            print(f"===> Bank signatures updated to: {new_text}")
            break


# =============================================================================
#  Spending threshold replacement (Sec 10.1 / 10.2)
# =============================================================================

def apply_spending_threshold(doc, data):
    """Replace the spending threshold in the limitations section.
    The template uses '$5,000' or '$5,000.00' — replace globally within the
    relevant section."""
    threshold = data.get('spendingThreshold')
    if not threshold or threshold == 5000:
        return  # Default matches template

    old_amounts = ['$5,000.00', '$5,000']
    new_amount = format_currency(threshold)

    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        for old_amt in old_amounts:
            while old_amt in full_text:
                idx = full_text.find(old_amt)
                _replace_span_in_paragraph(paragraph, idx, idx + len(old_amt), new_amount)
                full_text = ''.join(run.text for run in paragraph.runs)

    print(f"===> Spending threshold updated to {new_amount}")


# =============================================================================
#  Distribution settings (Sec 5.1)
# =============================================================================

def apply_distribution_settings(doc, data):
    """Update distribution frequency and dividend declaration period."""
    frequency = data.get('distributionFrequency', 'quarterly')

    # Sec 5.1: "Dividends shall be declared on a quarterly basis"
    if frequency and frequency.lower() != 'quarterly':
        freq_map = {
            'semi-annual': 'semi-annual',
            'semiannual': 'semi-annual',
            'annual': 'annual',
            'board_discretion': 'as-needed, at the discretion of the Board of Directors',
        }
        new_freq = freq_map.get(frequency.lower(), frequency)
        old_text = "on a quarterly basis"
        new_text = f"on a {new_freq} basis" if new_freq != freq_map.get('board_discretion') else new_freq

        for paragraph in doc.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            if old_text in full_text:
                idx = full_text.find(old_text)
                _replace_span_in_paragraph(paragraph, idx, idx + len(old_text), new_text)
                print(f"===> Distribution frequency updated to: {new_freq}")
                break


# =============================================================================
#  ROFR offer period replacement (Sec 13.1)
# =============================================================================

def apply_rofr_period(doc, data):
    """Update the ROFR offer period (default 180 days)."""
    rofr_days = data.get('rofrOfferDays')
    if not rofr_days or rofr_days == 180:
        return  # Default matches template

    old_text = "180 calendar days"
    new_text = f"{rofr_days} calendar days"

    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        if old_text in full_text:
            idx = full_text.find(old_text)
            _replace_span_in_paragraph(paragraph, idx, idx + len(old_text), new_text)
            print(f"===> ROFR offer period updated to {rofr_days} days")
            # Don't break — may appear in multiple paragraphs (13.1b and 13.1c)


# =============================================================================
#  Conditional section removal
# =============================================================================

def _find_section_paragraphs(doc, start_marker, end_markers):
    """Find all paragraphs belonging to a section.
    start_marker: text that begins the section (e.g., '13.1' or 'Right of First Refusal')
    end_markers: list of texts that signal the next section (e.g., ['13.2', '13.3', '14.'])
    Returns list of paragraph elements to remove.
    """
    paragraphs_to_remove = []
    in_section = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if not in_section:
            if start_marker in text:
                in_section = True
                paragraphs_to_remove.append(paragraph)
            continue

        # Check if we've hit the next section
        hit_end = False
        for end_marker in end_markers:
            if text.startswith(end_marker) or (end_marker in text and len(text) < 200):
                hit_end = True
                break

        if hit_end:
            break

        paragraphs_to_remove.append(paragraph)

    return paragraphs_to_remove


def _remove_paragraphs(doc, paragraphs):
    """Remove a list of paragraphs from the document."""
    removed = 0
    for p in paragraphs:
        parent = p._p.getparent()
        if parent is not None:
            parent.remove(p._p)
            removed += 1
    return removed


def _find_subsection_paragraphs(doc, parent_section_marker, subsection_marker, next_subsection_markers):
    """Find paragraphs of a subsection within a parent section.
    E.g., find (i) Drag Along within Section 13.3.
    """
    paragraphs_to_remove = []
    in_parent = False
    in_subsection = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if not in_parent:
            if parent_section_marker in text:
                in_parent = True
            continue

        if not in_subsection:
            if subsection_marker in text:
                in_subsection = True
                paragraphs_to_remove.append(paragraph)
            continue

        # Check if we've hit the next subsection or left the parent
        hit_end = False
        for marker in next_subsection_markers:
            if marker in text:
                hit_end = True
                break

        if hit_end:
            break

        paragraphs_to_remove.append(paragraph)

    return paragraphs_to_remove


def apply_conditional_removals(doc, data):
    """Remove sections based on boolean answers."""
    print("===> Applying conditional section removals...")

    # ROFR = No → remove Section 13.1 (Right of First Refusal)
    if data.get('rofr') is False:
        paras = _find_section_paragraphs(doc,
            'Right of First Refusal',
            ['13.2', '13.3', 'Purchase of Shareholder', '14.'])
        removed = _remove_paragraphs(doc, paras)
        print(f"===> ROFR=No: removed {removed} paragraphs (Sec 13.1)")

    # Drag-Along = No → remove 13.3(i)
    if data.get('dragAlong') is False:
        paras = _find_subsection_paragraphs(doc,
            '13.3', 'Drag Along',
            ['Tag Along', '(ii)', '14.'])
        removed = _remove_paragraphs(doc, paras)
        print(f"===> Drag-Along=No: removed {removed} paragraphs")

    # Tag-Along = No → remove 13.3(ii)
    if data.get('tagAlong') is False:
        paras = _find_subsection_paragraphs(doc,
            '13.3', 'Tag Along',
            ['14.', '15.', 'Withdrawing Shareholder'])
        removed = _remove_paragraphs(doc, paras)
        print(f"===> Tag-Along=No: removed {removed} paragraphs")

    # Non-compete = No → remove Sec 10.10 if it exists
    # Note: The SDD says non-compete is MISSING from the Corp template.
    # If it's been added, remove it when the answer is No.
    if data.get('nonCompete') is False:
        paras = _find_section_paragraphs(doc,
            '10.10', ['10.11', '11.', 'Voting Rights'])
        if paras:
            removed = _remove_paragraphs(doc, paras)
            print(f"===> Non-compete=No: removed {removed} paragraphs (Sec 10.10)")

    # Non-solicitation = No → remove Sec 10.9
    if data.get('nonSolicitation') is False:
        paras = _find_section_paragraphs(doc,
            '10.9', ['10.10', '11.', 'Voting Rights'])
        if paras:
            removed = _remove_paragraphs(doc, paras)
            print(f"===> Non-solicitation=No: removed {removed} paragraphs (Sec 10.9)")

    # Confidentiality/NDA = No → remove Sec 10.8
    if data.get('confidentiality') is False:
        paras = _find_section_paragraphs(doc,
            '10.8', ['10.9', '10.10', '11.', 'Voting Rights'])
        if paras:
            removed = _remove_paragraphs(doc, paras)
            print(f"===> Confidentiality=No: removed {removed} paragraphs (Sec 10.8)")


# =============================================================================
#  Dynamic shareholder handling
# =============================================================================

def _find_member_block(doc, member_index, context_text):
    """Find paragraphs containing placeholders for a specific member index
    near a context text (e.g., section heading).
    Returns list of paragraph objects."""
    num2 = f"{member_index:02d}"
    patterns = [
        f'shareholder_{num2}',
        f'shareholder_{member_index}',
        f'Shareholder_{num2}',
        f'Shareholder_{member_index}',
        f'Owner {member_index}',
    ]

    block = []
    in_context = False

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if context_text and context_text in text:
            in_context = True

        if in_context or not context_text:
            for pattern in patterns:
                if pattern in text:
                    block.append(paragraph)
                    break

    return block


def _clone_paragraph_after(doc, source_para, insert_after_para):
    """Clone a paragraph and insert it after another paragraph."""
    body = doc.element.body
    new_p = deepcopy(source_para._p)
    insert_after_para._p.addnext(new_p)
    return new_p


def handle_dynamic_shareholders(doc, data):
    """Handle 1-N shareholders in the document.
    Template has hardcoded shareholder blocks for shareholder_01 and shareholder_02.
    For 1 shareholder: remove shareholder_02 blocks.
    For 3+ shareholders: clone shareholder_02 blocks and update placeholders.
    """
    shareholders = data.get('members', []) or []
    if not shareholders:
        return

    num_shareholders = len(shareholders)
    print(f"===> Handling {num_shareholders} shareholders (template has 2)")

    if num_shareholders <= 2:
        # For 1 shareholder: find and remove all paragraphs with shareholder_02 placeholders
        if num_shareholders == 1:
            to_remove = []
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if ('shareholder_02' in text or 'shareholder_2' in text or
                    'Shareholder_02' in text or 'Shareholder_2' in text or
                    'member_02' in text or 'member_2' in text or
                    'Member_02' in text or 'Member_2' in text or
                    'Owner 2' in text):
                    to_remove.append(paragraph)
            removed = _remove_paragraphs(doc, to_remove)
            print(f"===> Removed {removed} paragraphs for absent shareholder_02")
        return

    # For 3+ shareholders: find shareholder_02 paragraphs, clone them for each additional
    # This is done in multiple passes for different sections (capital, MPI, signature)

    # Find paragraphs with shareholder_02 placeholders, grouped by proximity
    sh02_paragraphs = []
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if ('shareholder_02' in text or 'shareholder_2' in text or
            'Shareholder_02' in text or 'Shareholder_2' in text or
            'member_02' in text or 'member_2' in text or
            'Member_02' in text or 'Member_2' in text or
            'Owner 2' in text):
            sh02_paragraphs.append(paragraph)

    # Clone each sh02 paragraph for shareholders 3+
    for para in sh02_paragraphs:
        last_inserted = para
        for idx in range(3, num_shareholders + 1):
            new_p_element = _clone_paragraph_after(doc, para, last_inserted)
            # Update placeholders in the cloned paragraph
            # Replace 02/2 with the new index
            num2_new = f"{idx:02d}"
            for t_elem in new_p_element.iterchildren(qn('w:r')):
                for text_elem in t_elem.iterchildren(qn('w:t')):
                    if text_elem.text:
                        text_elem.text = (text_elem.text
                            .replace('shareholder_02', f'shareholder_{num2_new}')
                            .replace('shareholder_2', f'shareholder_{idx}')
                            .replace('Shareholder_02', f'Shareholder_{num2_new}')
                            .replace('Shareholder_2', f'Shareholder_{idx}')
                            .replace('member_02', f'member_{num2_new}')
                            .replace('member_2', f'member_{idx}')
                            .replace('Member_02', f'Member_{num2_new}')
                            .replace('Member_2', f'Member_{idx}')
                            .replace('Owner 2', f'Owner {idx}'))

            # Find the new paragraph wrapper for the next iteration
            # last_inserted needs to be a paragraph-like object with ._p
            class _ParagraphProxy:
                def __init__(self, p_element):
                    self._p = p_element
            last_inserted = _ParagraphProxy(new_p_element)

    print(f"===> Cloned shareholder_02 blocks for {num_shareholders - 2} additional shareholders")


# =============================================================================
#  Main placeholder replacement
# =============================================================================

def replace_placeholders(doc, data):
    """Replace all placeholders in the Shareholder Agreement document."""
    print("===> Replacing placeholders in Shareholder Agreement...")

    company_name = data.get('companyName', '')
    formation_state = data.get('formationState', '')
    formation_date_raw = data.get('formationDate', '') or data.get('paymentDate', '')
    company_address = data.get('companyAddress', '')
    county = data.get('county', '')
    total_shares = data.get('totalAuthorizedShares', '1,000')
    par_value = data.get('parValue', '0.01')

    formation_date = format_legal_date(formation_date_raw)
    witness_date = format_witness_date(formation_date_raw)

    # Members/shareholders
    members = data.get('members', []) or []
    managers = data.get('managers', []) or []
    directors = data.get('directors', []) or managers
    officers = data.get('officers', []) or []

    # Build replacement lists
    replacements = []  # (placeholder, value, bold)
    pct_placeholders = []  # (placeholder, pct_with_%, pct_without_%)

    def add(ph, val, bold=False):
        replacements.append((ph, val or '', bold))

    # --- Company info ---
    add('{{Company Name}}', company_name, True)
    add('{{COMPANY_NAME}}', company_name, True)
    add('{{Company Address}}', company_address)
    add('{{COMPANY_ADDRESS}}', company_address)
    add('{{Formation State}}', formation_state)
    add('{{FORMATION_STATE}}', formation_state)
    add('{{Formation Date}}', formation_date)
    add('{{FORMATION_DATE}}', formation_date)
    add('{{County}}', county)
    add('{{COUNTY}}', county)
    add('{{Total Authorized Shares}}', format_shares(total_shares))
    add('{{TOTAL_AUTHORIZED_SHARES}}', format_shares(total_shares))
    add('{{Par Value}}', format_currency(par_value))
    add('{{PAR_VALUE}}', format_currency(par_value))

    # --- Shareholders ---
    for idx, member in enumerate(members, start=1):
        num2 = f"{idx:02d}"
        name = member.get('name', '')
        shares = member.get('shares')
        capital = member.get('capitalContribution')
        ownership_pct = member.get('ownershipPercent', 0)

        if ownership_pct is None:
            ownership_pct = 0

        pct_str = format_percentage(ownership_pct)
        pct_str_no_pct = pct_str.rstrip('%')

        # Name placeholders (all variants)
        for ph in (
            f'{{{{shareholder_{num2}_full_name}}}}',
            f'{{{{shareholder_{idx}_full_name}}}}',
            f'{{{{Shareholder_{num2}_full_name}}}}',
            f'{{{{Shareholder_{idx}_full_name}}}}',
            f'{{{{member_{num2}_full_name}}}}',
            f'{{{{member_{idx}_full_name}}}}',
            f'{{{{Member_{num2}_full_name}}}}',
            f'{{{{Member_{idx}_full_name}}}}',
            f'{{{{Owner {idx} Name}}}}',
        ):
            add(ph, name, True)

        # Shares
        if shares is not None:
            shares_str = format_shares(shares)
            for ph in (
                f'{{{{shareholder_{num2}_shares}}}}',
                f'{{{{shareholder_{idx}_shares}}}}',
                f'{{{{Shareholder_{num2}_shares}}}}',
                f'{{{{Shareholder_{idx}_shares}}}}',
                f'{{{{Owner {idx} Ownership #Shares}}}}',
            ):
                add(ph, shares_str)

        # Capital contribution
        if capital is not None:
            capital_str = format_currency(capital)
            for ph in (
                f'{{{{shareholder_{num2}_capital}}}}',
                f'{{{{shareholder_{idx}_capital}}}}',
                f'{{{{Shareholder_{num2}_capital}}}}',
                f'{{{{Shareholder_{idx}_capital}}}}',
                f'{{{{Owner {idx} Capital}}}}',
            ):
                add(ph, capital_str)

        # Ownership percentage
        for ph in (
            f'{{{{shareholder_{num2}_pct}}}}',
            f'{{{{shareholder_{idx}_pct}}}}',
            f'{{{{Shareholder_{num2}_pct}}}}',
            f'{{{{Shareholder_{idx}_pct}}}}',
            f'{{{{member_{num2}_pct}}}}',
            f'{{{{member_{idx}_pct}}}}',
            f'{{{{Owner {idx} Ownership %}}}}',
        ):
            pct_placeholders.append((ph, pct_str, pct_str_no_pct))

    # --- Directors ---
    for idx, director in enumerate(directors, start=1):
        num2 = f"{idx:02d}"
        director_name = director.get('name', '')
        for ph in (
            f'{{{{Director_{num2}_Name}}}}',
            f'{{{{Director_{idx}_Name}}}}',
            f'{{{{Director {idx} Name}}}}',
        ):
            add(ph, director_name, True)

    # --- Officers ---
    for idx, officer in enumerate(officers, start=1):
        num2 = f"{idx:02d}"
        officer_name = officer.get('name', '')
        officer_role = officer.get('role', '') or officer.get('title', '')
        for ph in (
            f'{{{{Officer_{num2}_Name}}}}',
            f'{{{{Officer_{idx}_Name}}}}',
            f'{{{{Officer {idx} Name}}}}',
        ):
            add(ph, officer_name, True)
        for ph in (
            f'{{{{Officer_{num2}_Role}}}}',
            f'{{{{Officer_{idx}_Role}}}}',
            f'{{{{Officer {idx} Role}}}}',
        ):
            add(ph, officer_role)

    # --- Process all paragraphs ---
    WITNESS_DATE_PLACEHOLDERS = ('{{Formation Date}}', '{{FORMATION_DATE}}',
                                  '{{Payment Date}}')

    def process_paragraph(paragraph, in_witness_block=False):
        if '{{' not in paragraph.text:
            return

        modified_runs = []
        for placeholder, value, bold in replacements:
            # Use witness date format in the IN WITNESS WHEREOF block
            if in_witness_block and placeholder in WITNESS_DATE_PLACEHOLDERS:
                value = witness_date
            modified_runs.extend(
                _replace_placeholder_in_paragraph(paragraph, placeholder, value, bold)
            )

        for placeholder, pct_str, pct_no_pct in pct_placeholders:
            _replace_pct_placeholder_in_paragraph(paragraph, placeholder, pct_str, pct_no_pct)

        for run in modified_runs:
            _enforce_font(run)

    # Body paragraphs
    in_witness = False
    for paragraph in doc.paragraphs:
        if 'IN WITNESS WHEREOF' in paragraph.text:
            in_witness = True
        process_paragraph(paragraph, in_witness_block=in_witness)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    print("===> Placeholder replacement complete")


# =============================================================================
#  Post-processing
# =============================================================================

def _add_keep_next(paragraph):
    """Add keepNext property to a paragraph so it stays with the next paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:keepNext')):
        pPr.remove(existing)
    keep_next = OxmlElement('w:keepNext')
    pPr.append(keep_next)


def _is_section_heading(text):
    """Detect section headings in the Shareholder Agreement template."""
    stripped = text.strip()
    if not stripped:
        return False
    # ARTICLE headings
    if re.match(r'^ARTICLE\s+[IVXLC]+', stripped):
        return True
    # Numbered sections: "4.2 Initial Capital" etc.
    if re.match(r'^\d+\.\d*\s+[A-Z]', stripped):
        return True
    # All-caps headings
    if len(stripped) > 5 and stripped == stripped.upper() and re.match(r'^[A-Z\s\-–,\.]+$', stripped):
        return True
    return False


def post_process_shareholder_agreement(doc):
    """Best-practice formatting fixes for Shareholder Agreement documents."""
    print("===> Post-processing shareholder agreement...")

    paragraphs = doc.paragraphs

    # --- 1. keepNext on section headings ---
    headings_fixed = 0
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if _is_section_heading(text):
            _add_keep_next(paragraph)
            headings_fixed += 1
            if i + 1 < len(paragraphs) and not paragraphs[i + 1].text.strip():
                _add_keep_next(paragraphs[i + 1])

    # --- 2. [SIGNATURE PAGE BELOW] + pageBreakBefore on IN WITNESS WHEREOF ---
    witness_para = None
    for paragraph in doc.paragraphs:
        if 'IN WITNESS WHEREOF' in paragraph.text:
            witness_para = paragraph
            break

    if witness_para is not None:
        body = witness_para._p.getparent()

        # Remove trailing empties before witness
        empties_removed = 0
        while True:
            prev_el = witness_para._p.getprevious()
            if prev_el is None:
                break
            prev_txts = []
            for r_el in prev_el.iterchildren(qn('w:r')):
                for t_el in r_el.iterchildren(qn('w:t')):
                    prev_txts.append(t_el.text or '')
            if ''.join(prev_txts).strip() == '':
                body.remove(prev_el)
                empties_removed += 1
            else:
                break

        # Check if [SIGNATURE PAGE BELOW] already exists
        prev = witness_para._p.getprevious()
        prev_text = ''
        if prev is not None:
            prev_txts = []
            for r_el in prev.iterchildren(qn('w:r')):
                for t_el in r_el.iterchildren(qn('w:t')):
                    prev_txts.append(t_el.text or '')
            prev_text = ''.join(prev_txts).strip()

        if 'SIGNATURE PAGE BELOW' not in prev_text:
            sig_p = OxmlElement('w:p')
            sig_pPr = OxmlElement('w:pPr')
            sig_jc = OxmlElement('w:jc')
            sig_jc.set(qn('w:val'), 'center')
            sig_pPr.append(sig_jc)
            sig_p.append(sig_pPr)
            sig_r = OxmlElement('w:r')
            sig_rPr = OxmlElement('w:rPr')
            sig_rFonts = OxmlElement('w:rFonts')
            sig_rFonts.set(qn('w:ascii'), 'Times New Roman')
            sig_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            sig_rPr.append(sig_rFonts)
            sig_sz = OxmlElement('w:sz')
            sig_sz.set(qn('w:val'), '24')  # 12pt
            sig_rPr.append(sig_sz)
            sig_r.append(sig_rPr)
            sig_t = OxmlElement('w:t')
            sig_t.text = '[SIGNATURE PAGE BELOW]'
            sig_r.append(sig_t)
            sig_p.append(sig_r)
            body.insert(list(body).index(witness_para._p), sig_p)
            print("===> Inserted [SIGNATURE PAGE BELOW]")

        # Ensure pageBreakBefore on WITNESS WHEREOF
        pPr = witness_para._p.get_or_add_pPr()
        if not pPr.findall(qn('w:pageBreakBefore')):
            page_break = OxmlElement('w:pageBreakBefore')
            pPr.append(page_break)
            print("===> Added pageBreakBefore to IN WITNESS WHEREOF")

    # --- 3. Remove excess empty paragraphs in signature section ---
    in_sig = False
    to_remove = []
    consecutive = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if 'IN WITNESS WHEREOF' in text:
            in_sig = True
            consecutive = 0
            continue
        if in_sig:
            if not text:
                consecutive += 1
                if consecutive > 2:
                    to_remove.append(paragraph)
            else:
                consecutive = 0
    sig_removed = _remove_paragraphs(doc, to_remove)

    # --- 4. Remove "PAGE X" footer text ---
    pages_removed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r'^PAGE\s+\d+$', text):
            parent = paragraph._p.getparent()
            if parent is not None:
                parent.remove(paragraph._p)
                pages_removed += 1

    # --- 5. Remove trailing empty paragraphs ---
    body_el = doc.element.body
    trailing_removed = 0
    while len(doc.paragraphs) > 0:
        last_p = doc.paragraphs[-1]
        if not last_p.text.strip():
            body_el.remove(last_p._p)
            trailing_removed += 1
        else:
            break

    # --- 6. Normalize signature block lines ---
    SIG_LINE_TAB_COUNT = '\t\t\t\t\t\t'
    in_sig = False
    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        if 'WITNESS WHEREOF' in txt:
            in_sig = True
        if not in_sig or not paragraph.runs:
            continue
        first_run = paragraph.runs[0]
        needs_norm = False
        if re.match(r'^\d+(\.\d+)?%?\s*Owner', txt):
            needs_norm = True
        elif re.match(r'^By:\s*_', txt):
            needs_norm = True
        elif re.match(r'^Name:\s', txt):
            needs_norm = True
        elif re.match(r'^Title:\s', txt):
            needs_norm = True
        elif txt in ('SHAREHOLDER', 'SHAREHOLDERS', 'CORPORATION'):
            needs_norm = True
        if needs_norm:
            raw = first_run.text
            stripped = raw.lstrip()
            if stripped:
                first_run.text = SIG_LINE_TAB_COUNT + stripped
            else:
                first_run.text = SIG_LINE_TAB_COUNT

    print(f"===> Post-processing done: {headings_fixed} headings keepNext, "
          f"{sig_removed} sig empties, {pages_removed} PAGE X removed, "
          f"{trailing_removed} trailing empties removed")


# =============================================================================
#  Lambda handler
# =============================================================================

def lambda_handler(event, context):
    print("===> Shareholder Agreement Lambda invoked")
    print("===> RAW EVENT:")
    print(json.dumps(event))

    try:
        if "body" in event:
            body = json.loads(event["body"])
        else:
            body = event

        form_data = body.get("form_data")
        s3_bucket = body.get("s3_bucket", OUTPUT_BUCKET)
        s3_key = body.get("s3_key")
        template_url = body.get("templateUrl")
        return_docx = body.get("return_docx", False)

        if not form_data:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Missing 'form_data'"})
            }
        if not s3_bucket or not s3_key:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Missing 's3_bucket' or 's3_key'"})
            }
        if not template_url:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Missing 'templateUrl'"})
            }

        template_bucket, template_key = extract_s3_info(template_url)
        if not template_bucket or not template_key:
            template_bucket = BUCKET_NAME
            template_key = 'templates/shareholder-agreement-template.docx'
            print(f"===> Could not parse templateUrl, using default: "
                  f"s3://{template_bucket}/{template_key}")
        else:
            print(f"===> Template: s3://{template_bucket}/{template_key}")

        print(f"===> Output: s3://{s3_bucket}/{s3_key}")

    except Exception as e:
        return {
            "statusCode": 400,
            "body": json.dumps({"error": "Invalid input payload", "details": str(e)})
        }

    tmpdir = tempfile.gettempdir()
    template_path = os.path.join(tmpdir, "template_shareholder_agreement.docx")
    output_path = os.path.join(tmpdir, "filled_shareholder_agreement.docx")

    try:
        # Download template from S3
        download_from_s3(template_bucket, template_key, template_path)

        # Load Word document
        print("===> Loading Word document...")
        doc = Document(template_path)

        # 1. Handle dynamic shareholders (clone/remove blocks BEFORE placeholder replacement)
        handle_dynamic_shareholders(doc, form_data)

        # 2. Replace all {{placeholders}} with form data values
        replace_placeholders(doc, form_data)

        # 3. Update Majority/Super Majority definitions (Sec 1.6 / 1.11)
        apply_majority_definition(doc, form_data)

        # 4. Apply section-specific voting text replacements
        apply_voting_replacements(doc, form_data)

        # 5. Bank signature replacement (Sec 10.7)
        apply_bank_signature_replacement(doc, form_data)

        # 6. Spending threshold replacement
        apply_spending_threshold(doc, form_data)

        # 7. Distribution settings (frequency, dividends)
        apply_distribution_settings(doc, form_data)

        # 8. ROFR offer period
        apply_rofr_period(doc, form_data)

        # 9. Conditional section removal (ROFR, Drag/Tag-Along, Non-compete, etc.)
        apply_conditional_removals(doc, form_data)

        # 10. Post-processing (formatting fixes)
        post_process_shareholder_agreement(doc)

        # Save filled document
        print("===> Saving filled document...")
        doc.save(output_path)
        print(f"===> Saved to {output_path}")

        # Upload to S3
        upload_to_s3(output_path, s3_bucket, s3_key)

        # Return response
        if return_docx:
            with open(output_path, 'rb') as f:
                docx_content = f.read()
            encoded = base64.b64encode(docx_content).decode('utf-8')
            return {
                "statusCode": 200,
                "headers": {
                    "Content-Type": "application/json"
                },
                "body": json.dumps({
                    "message": "Shareholder Agreement generated successfully",
                    "docx_base64": encoded
                })
            }
        else:
            return {
                "statusCode": 200,
                "body": json.dumps({
                    "message": "Shareholder Agreement generated successfully",
                    "s3_bucket": s3_bucket,
                    "s3_key": s3_key,
                    "s3_url": f"s3://{s3_bucket}/{s3_key}"
                })
            }

    except Exception as e:
        import traceback
        error_msg = str(e)
        traceback.print_exc()
        print(f"===> ERROR: {error_msg}")
        return {
            "statusCode": 500,
            "body": json.dumps({
                "error": "Failed to generate shareholder agreement",
                "details": error_msg
            })
        }
