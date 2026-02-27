import os
import json
import tempfile
import boto3
import re
import base64
from copy import deepcopy
from datetime import datetime
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Constants
TEMPLATE_BUCKET = os.environ.get('TEMPLATE_BUCKET', 'avenida-legal-documents')
OUTPUT_BUCKET = os.environ.get('OUTPUT_BUCKET', 'avenida-legal-documents')
BUCKET_NAME = os.environ.get('BUCKET_NAME', TEMPLATE_BUCKET)

# Initialize S3 client
s3_client = boto3.client('s3', region_name=os.environ.get('AWS_REGION', 'us-west-1'))


def download_from_s3(bucket, key, local_path):
    """Download file from S3 to local path"""
    print(f"===> Downloading s3://{bucket}/{key} to {local_path}")
    s3_client.download_file(bucket, key, local_path)
    print("===> Downloaded successfully")


def upload_to_s3(local_path, bucket, key):
    """Upload file from local path to S3"""
    print(f"===> Uploading {local_path} to s3://{bucket}/{key}")
    s3_client.upload_file(local_path, bucket, key)
    print("===> Uploaded successfully")


def extract_s3_info(url):
    """Extract bucket and key from S3 URL"""
    if not url:
        return None, None

    if url.startswith('s3://'):
        parts = url[5:].split('/', 1)
        if len(parts) == 2:
            return parts[0], parts[1]
        return parts[0], ''

    if url.startswith('https://'):
        match = re.match(r'https://([^.]+)\.s3[^/]*\.amazonaws\.com/(.+)', url)
        if match:
            return match.group(1), match.group(2)

        match = re.match(r'https://s3[^/]*\.amazonaws\.com/([^/]+)/(.+)', url)
        if match:
            return match.group(1), match.group(2)

    return None, None


def _format_witness_date(value: str) -> str:
    """Format date as numeric ordinal for IN WITNESS WHEREOF block: e.g. 8th day of February, 2026."""
    if not value or not value.strip():
        return value
    # Already in "Xth day of Month, year" form
    if 'day of' in value:
        m = re.match(r'^\s*(\d{1,2})(st|nd|rd|th)\s+day of\s+.+', value, re.IGNORECASE)
        if m:
            return value.strip()
    # Parse m/d/yyyy or yyyy-mm-dd
    match = re.match(r'^\s*(\d{1,2})/(\d{1,2})/(\d{4})\s*$', value) or re.match(r'^\s*(\d{4})-(\d{1,2})-(\d{1,2})\s*$', value)
    if not match:
        return value
    if len(match.group(1)) == 4:
        year, month, day = int(match.group(1)), int(match.group(2)), int(match.group(3))
    else:
        month, day, year = int(match.group(1)), int(match.group(2)), int(match.group(3))
    try:
        dt = datetime(year, month, day)
    except ValueError:
        return value
    month_name = dt.strftime('%B')
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return f"{day}{suffix} day of {month_name}, {year}"


# ---------- Run-level replacement (non-destructive) ----------

def _insert_run_after(paragraph, run, text):
    """Insert a new run after the given run, returning the new run."""
    new_run = paragraph.add_run(text)
    run._element.addnext(new_run._element)
    return new_run


def _apply_run_format(source_run, target_run):
    """Copy formatting from source_run to target_run, preserving font properties."""
    if source_run._element.rPr is not None:
        target_run._element.insert(0, deepcopy(source_run._element.rPr))


def _enforce_font(run):
    """Ensure a modified run uses 12pt Times New Roman."""
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def _replace_span_in_paragraph(paragraph, start, end, value):
    """Replace a character span [start, end) in the paragraph's runs with value.
    Preserves run formatting. Returns the run containing the replacement value, or None."""
    pos = 0
    run_start_idx = None
    run_end_idx = None
    run_start_offset = 0
    run_end_offset = 0

    for i, run in enumerate(paragraph.runs):
        run_len = len(run.text)
        if run_start_idx is None and pos + run_len > start:
            run_start_idx = i
            run_start_offset = start - pos
        if run_start_idx is not None and pos + run_len >= end:
            run_end_idx = i
            run_end_offset = end - pos
            break
        pos += run_len

    if run_start_idx is None or run_end_idx is None:
        return None

    if run_start_idx == run_end_idx:
        run = paragraph.runs[run_start_idx]
        prefix = run.text[:run_start_offset]
        suffix = run.text[run_end_offset:]

        if prefix:
            run.text = prefix
            value_run = _insert_run_after(paragraph, run, value)
            _apply_run_format(run, value_run)
        else:
            run.text = value
            value_run = run

        if suffix:
            suffix_run = _insert_run_after(paragraph, value_run, suffix)
            _apply_run_format(run, suffix_run)

        return value_run
    else:
        start_run = paragraph.runs[run_start_idx]
        end_run = paragraph.runs[run_end_idx]
        prefix = start_run.text[:run_start_offset]
        suffix = end_run.text[run_end_offset:]

        # Clear middle runs
        for i in range(run_start_idx + 1, run_end_idx):
            paragraph.runs[i].text = ''

        if prefix:
            start_run.text = prefix
            value_run = _insert_run_after(paragraph, start_run, value)
            _apply_run_format(start_run, value_run)
        else:
            start_run.text = value
            value_run = start_run

        if suffix:
            end_run.text = suffix
        else:
            end_run.text = ''

        return value_run


def _replace_placeholder_in_paragraph(paragraph, placeholder, value):
    """Find all occurrences of placeholder in paragraph and replace with value (run-level).
    Returns list of runs that were modified (for font enforcement)."""
    if placeholder not in paragraph.text:
        return []

    modified_runs = []
    full_text = ''.join(run.text for run in paragraph.runs)
    search_start = 0
    while True:
        idx = full_text.find(placeholder, search_start)
        if idx == -1:
            break
        end = idx + len(placeholder)
        result_run = _replace_span_in_paragraph(paragraph, idx, end, value)
        if result_run is None:
            break
        modified_runs.append(result_run)
        full_text = ''.join(run.text for run in paragraph.runs)
        search_start = idx + len(value)
    return modified_runs


def replace_placeholders(doc, data):
    """Replace placeholders in Bylaws document using run-level replacement.
    IN WITNESS WHEREOF block gets numeric ordinal date.
    Enforces 12pt Times New Roman on all replaced values."""
    print("===> Replacing placeholders in document...")

    company_name = data.get('companyName', '')
    formation_state = data.get('formationState', '')
    payment_date = data.get('paymentDate', '')
    witness_date = _format_witness_date(payment_date)
    number_of_shares = str(data.get('numberOfShares', '') or '')
    officer_1_name = data.get('officer1Name', '')
    officer_1_role = data.get('officer1Role', '')

    # Build owner names (1-6)
    owner_names = {}
    for i in range(1, 7):
        owner_names[i] = data.get(f'owner{i}Name', '')

    # Map of placeholder → value (non-date)
    placeholders = {
        '{{Company Name}}': company_name,
        '{{Formation State}}': formation_state,
        '{{Number of Shares}}': number_of_shares,
        '{{Officer 1 Name}}': officer_1_name,
        '{{Officer 1 Role}}': officer_1_role,
    }
    for i in range(1, 7):
        placeholders[f'{{{{Owner {i} Name}}}}'] = owner_names[i]

    def process_paragraph(paragraph):
        full_text = paragraph.text
        if '{{' not in full_text:
            return

        in_witness = "IN WITNESS WHEREOF" in full_text
        modified_runs = []

        # Replace date placeholder with appropriate format
        date_value = witness_date if in_witness else payment_date
        modified_runs.extend(
            _replace_placeholder_in_paragraph(paragraph, '{{Payment Date}}', date_value)
        )

        # Replace all other placeholders
        for ph, val in placeholders.items():
            modified_runs.extend(
                _replace_placeholder_in_paragraph(paragraph, ph, val)
            )

        # Enforce font on all modified runs
        for run in modified_runs:
            _enforce_font(run)

    # Process body paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)


# ---------- Post-processing: fix template formatting issues ----------

def _is_heading_paragraph(text):
    """Detect section headings in the Bylaws template.
    Matches: 'ARTICLE I – OFFICES', '1.    PLACE OF MEETINGS', 'SHAREHOLDERS', etc."""
    stripped = text.strip()
    if not stripped:
        return False
    # ARTICLE headings: "ARTICLE I – OFFICES", "ARTICLE XV - EMERGENCY BY-LAWS"
    if re.match(r'^ARTICLE\s+[IVXLC]+', stripped):
        return True
    # Numbered sections: "1.    PLACE OF MEETINGS", "10.  NOTICE OF MEETINGS"
    if re.match(r'^\d+\.\s{2,}[A-Z]', stripped):
        return True
    # All-caps standalone headings (>5 chars, no lowercase)
    if len(stripped) > 5 and stripped == stripped.upper() and re.match(r'^[A-Z\s\-–,\.]+$', stripped):
        return True
    return False


def _add_keep_next(paragraph):
    """Add keepNext property to a paragraph so it stays with the next paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    # Remove existing keepNext if any
    for existing in pPr.findall(qn('w:keepNext')):
        pPr.remove(existing)
    keep_next = OxmlElement('w:keepNext')
    pPr.append(keep_next)


def post_process_bylaws(doc):
    """Fix template formatting issues after placeholder replacement:
    1. Add keepNext to all section headings AND the empty paragraph following them
       (the template has heading → empty para → body text; we need the chain to hold)
    2. Fix 'theretofore' typo → 'therefore'
    3. Remove excess empty paragraphs in signature section (keep max 1 between blocks)
    """
    print("===> Post-processing: fixing template formatting...")

    paragraphs = doc.paragraphs
    headings_fixed = 0
    typos_fixed = 0

    # --- 1. Add keepNext to heading paragraphs AND the empty para that follows ---
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if _is_heading_paragraph(text):
            _add_keep_next(paragraph)
            headings_fixed += 1
            # Also add keepNext to the next paragraph if it is empty (spacer between
            # heading and body text).  This ensures heading + spacer + body stay together.
            if i + 1 < len(paragraphs) and not paragraphs[i + 1].text.strip():
                _add_keep_next(paragraphs[i + 1])

    # --- 2. Fix "theretofore" → "therefore" typo ---
    for paragraph in paragraphs:
        if 'theretofore' in paragraph.text.lower():
            for run in paragraph.runs:
                if 'theretofore' in run.text:
                    run.text = run.text.replace('theretofore', 'therefore')
                    typos_fixed += 1
                if 'Theretofore' in run.text:
                    run.text = run.text.replace('Theretofore', 'Therefore')
                    typos_fixed += 1

    # --- 3. Insert "[SIGNATURE PAGE BELOW]" + pageBreakBefore on WITNESS WHEREOF ---
    # The Org Resolution template already has this, but the Bylaws template does not.
    # "IN WITNESS WHEREOF" must always be the first line of the last page, with
    # "[SIGNATURE PAGE BELOW]" as the last text on the preceding page.
    witness_para = None
    for paragraph in doc.paragraphs:
        if 'IN WITNESS WHEREOF' in paragraph.text:
            witness_para = paragraph
            break

    if witness_para is not None:
        body = witness_para._p.getparent()

        # Check if "[SIGNATURE PAGE BELOW]" already exists before the witness paragraph
        prev = witness_para._p.getprevious()
        prev_text = ''
        if prev is not None:
            prev_texts = []
            for t in prev.iterchildren(qn('w:r')):
                for tt in t.iterchildren(qn('w:t')):
                    prev_texts.append(tt.text or '')
            prev_text = ''.join(prev_texts).strip()

        # Remove trailing empty paragraphs immediately before witness (or before
        # the [SIGNATURE PAGE BELOW] we are about to insert) — these just add whitespace.
        empties_removed_before = 0
        while True:
            prev_el = witness_para._p.getprevious()
            if prev_el is None:
                break
            prev_txts = []
            for r_el in prev_el.iterchildren(qn('w:r')):
                for t_el in r_el.iterchildren(qn('w:t')):
                    prev_txts.append(t_el.text or '')
            if ''.join(prev_txts).strip() == '':
                body.remove(prev_el)
                empties_removed_before += 1
            else:
                break
        if empties_removed_before:
            print(f"===> Removed {empties_removed_before} empty paragraphs before signature page")

        if 'SIGNATURE PAGE BELOW' not in prev_text:
            # Insert a new centered "[SIGNATURE PAGE BELOW]" paragraph before witness
            sig_p = OxmlElement('w:p')
            sig_pPr = OxmlElement('w:pPr')
            sig_jc = OxmlElement('w:jc')
            sig_jc.set(qn('w:val'), 'center')
            sig_pPr.append(sig_jc)
            sig_p.append(sig_pPr)
            sig_r = OxmlElement('w:r')
            # Copy font formatting from witness paragraph
            sig_rPr = OxmlElement('w:rPr')
            sig_rFonts = OxmlElement('w:rFonts')
            sig_rFonts.set(qn('w:ascii'), 'Times New Roman')
            sig_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            sig_rPr.append(sig_rFonts)
            sig_sz = OxmlElement('w:sz')
            sig_sz.set(qn('w:val'), '24')  # 12pt
            sig_rPr.append(sig_sz)
            sig_r.append(sig_rPr)
            sig_t = OxmlElement('w:t')
            sig_t.text = '[SIGNATURE PAGE BELOW]'
            sig_r.append(sig_t)
            sig_p.append(sig_r)
            body.insert(list(body).index(witness_para._p), sig_p)
            print("===> Inserted [SIGNATURE PAGE BELOW] paragraph")

        # Add pageBreakBefore to "IN WITNESS WHEREOF" paragraph
        pPr = witness_para._p.get_or_add_pPr()
        # Remove existing pageBreakBefore if any
        for existing in pPr.findall(qn('w:pageBreakBefore')):
            pPr.remove(existing)
        page_break = OxmlElement('w:pageBreakBefore')
        pPr.append(page_break)
        print("===> Added pageBreakBefore to IN WITNESS WHEREOF")

    # --- 4. Remove excess empty paragraphs in signature section ---
    # The template has 3-4 empty paragraphs between each signature block.
    # Keep at most 2 empty paragraphs between blocks for comfortable spacing.
    in_signature_section = False
    paragraphs_to_remove = []
    consecutive_empties = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if 'IN WITNESS WHEREOF' in text:
            in_signature_section = True
            consecutive_empties = 0
            continue

        if in_signature_section:
            if not text:
                consecutive_empties += 1
                if consecutive_empties > 2:
                    # Mark for removal — keep 2 empty paras between blocks
                    paragraphs_to_remove.append(paragraph)
            else:
                consecutive_empties = 0

    # Remove the excess empty paragraphs by deleting their XML elements
    removed = 0
    for paragraph in paragraphs_to_remove:
        p_element = paragraph._p
        parent = p_element.getparent()
        if parent is not None:
            parent.remove(p_element)
            removed += 1

    # --- 5. Remove "PAGE X" footer text from signature section ---
    pages_removed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r'^PAGE\s+\d+$', text):
            p_element = paragraph._p
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)
                pages_removed += 1

    # --- 6. Remove trailing empty paragraphs at end of document ---
    body_el = doc.element.body
    trailing_removed = 0
    while len(doc.paragraphs) > 0:
        last_p = doc.paragraphs[-1]
        if not last_p.text.strip():
            body_el.remove(last_p._p)
            trailing_removed += 1
        else:
            break

    # --- 7. Normalize lettered sub-item indentation (A., B., C., D., E.) ---
    # Some template sub-items have extra left indentation (720 twips) while
    # others have 0.  Normalize all lettered sub-items to left=0.
    indent_fixes = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r'^[A-Z]\.\s', text):
            pPr = paragraph._p.get_or_add_pPr()
            ind_elements = pPr.findall(qn('w:ind'))
            for ind_el in ind_elements:
                current_left = ind_el.get(qn('w:left'))
                if current_left and int(current_left) > 0:
                    ind_el.set(qn('w:left'), '0')
                    indent_fixes += 1

    print(f"===> Post-processing done: {headings_fixed} headings keepNext, {typos_fixed} typos fixed, {indent_fixes} indents normalized, {removed} excess empty paras removed from signature")


def lambda_handler(event, context):
    print("===> Bylaws Lambda invoked")

    try:
        body = json.loads(event.get('body', '{}'))
    except Exception:
        body = event if isinstance(event, dict) else {}

    form_data = body.get('form_data', {}) or {}
    s3_bucket = body.get('s3_bucket', OUTPUT_BUCKET)
    s3_key = body.get('s3_key', '')
    template_url = body.get('templateUrl')
    return_docx = body.get('return_docx', False)

    if not template_url:
        return {
            "statusCode": 400,
            "body": json.dumps({"error": "Missing 'templateUrl'"})
        }

    template_bucket, template_key = extract_s3_info(template_url)
    if not template_bucket or not template_key:
        template_bucket = BUCKET_NAME
        template_key = 'templates/bylaws-template.docx'
        print(f"===> Could not parse templateUrl, using default: s3://{template_bucket}/{template_key}")

    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "template_bylaws.docx")
        output_path = os.path.join(tmpdir, "filled_bylaws.docx")

        download_from_s3(template_bucket, template_key, template_path)

        doc = Document(template_path)
        replace_placeholders(doc, form_data)
        post_process_bylaws(doc)
        doc.save(output_path)

        upload_to_s3(output_path, s3_bucket, s3_key)

        if return_docx:
            with open(output_path, "rb") as f:
                docx_content = f.read()
            encoded = base64.b64encode(docx_content).decode('utf-8')
            return {
                "statusCode": 200,
                "headers": {"Content-Type": "application/json"},
                "body": json.dumps({
                    "message": "Bylaws generated successfully",
                    "docx_base64": encoded
                })
            }

        return {
            "statusCode": 200,
            "body": json.dumps({"message": "Bylaws generated successfully"})
        }

    return {
        "statusCode": 500,
        "body": json.dumps({"error": "Failed to generate bylaws"})
    }
