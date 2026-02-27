import os
import json
import tempfile
import boto3
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import base64
from copy import deepcopy
from datetime import datetime

# Constants
TEMPLATE_BUCKET = os.environ.get('TEMPLATE_BUCKET', 'company-formation-template-llc-and-inc')
OUTPUT_BUCKET = os.environ.get('OUTPUT_BUCKET', 'avenida-legal-documents')
BUCKET_NAME = os.environ.get('BUCKET_NAME', TEMPLATE_BUCKET)

# Initialize S3 client
s3_client = boto3.client('s3', region_name='us-west-1')

def download_from_s3(bucket, key, local_path):
    """Download file from S3 to local path"""
    print(f"===> Downloading s3://{bucket}/{key} to {local_path}")
    s3_client.download_file(bucket, key, local_path)
    print(f"===> Downloaded successfully")

def upload_to_s3(local_path, bucket, key):
    """Upload file from local path to S3"""
    print(f"===> Uploading {local_path} to s3://{bucket}/{key}")
    s3_client.upload_file(local_path, bucket, key)
    print(f"===> Uploaded successfully")

def extract_s3_info(url):
    """Extract bucket and key from S3 URL"""
    if not url:
        return None, None
    
    if url.startswith('s3://'):
        parts = url[5:].split('/', 1)
        if len(parts) == 2:
            return parts[0], parts[1]
        return parts[0], ''
    
    if url.startswith('https://'):
        match = re.match(r'https://([^.]+)\.s3[^/]*\.amazonaws\.com/(.+)', url)
        if match:
            return match.group(1), match.group(2)
        
        match = re.match(r'https://s3[^/]*\.amazonaws\.com/([^/]+)/(.+)', url)
        if match:
            return match.group(1), match.group(2)
    
    return None, None

def format_address(address_str):
    """Format address string, handling None and empty values"""
    if not address_str or address_str.strip() == '':
        return ''
    return address_str.strip()

def format_percentage(value):
    """Format ownership percentage - pixel perfect formatting (no trailing zeros)"""
    if value is None:
        return "0%"
    
    if isinstance(value, (int, float)):
        num = float(value)
    elif isinstance(value, str):
        try:
            num = float(value.replace(',', '.'))
        except:
            return "0%"
    else:
        return "0%"
    
    if 0 <= num <= 1:
        num = num * 100
    
    if num == int(num):
        return f"{int(num)}%"
    else:
        formatted = f"{num:.10f}".rstrip('0').rstrip('.')
        return f"{formatted}%"

def replace_placeholders(doc, data):
    """Replace placeholders in Organizational Resolution document"""
    print("===> Replacing placeholders in document...")

    # Company information
    company_name = data.get('companyName', '')
    company_address = format_address(data.get('companyAddress', ''))
    formation_state = data.get('formationState', '')
    # For body: use formatted formationDate; for witness date, prefer raw paymentDateRaw fallback to formationDate
    formation_date_raw = data.get('formationDate', '')
    payment_date_raw = data.get('paymentDateRaw', '') or formation_date_raw

    def number_to_word(n: int) -> str:
        words = {
            1: 'first', 2: 'second', 3: 'third', 4: 'fourth', 5: 'fifth', 6: 'sixth', 7: 'seventh',
            8: 'eighth', 9: 'ninth', 10: 'tenth', 11: 'eleventh', 12: 'twelfth', 13: 'thirteenth',
            14: 'fourteenth', 15: 'fifteenth', 16: 'sixteenth', 17: 'seventeenth', 18: 'eighteenth',
            19: 'nineteenth', 20: 'twentieth', 21: 'twenty-first', 22: 'twenty-second', 23: 'twenty-third',
            24: 'twenty-fourth', 25: 'twenty-fifth', 26: 'twenty-sixth', 27: 'twenty-seventh',
            28: 'twenty-eighth', 29: 'twenty-ninth', 30: 'thirtieth', 31: 'thirty-first',
        }
        return words.get(n, str(n))

    def ordinal_suffix(day: int) -> str:
        if 11 <= day <= 13:
            return 'th'
        return {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')

    def _parse_date(value: str):
        """Parse m/d/yyyy or yyyy-mm-dd; return (day, month_name, year) or None."""
        if not value or not value.strip():
            return None
        match = re.match(r'^\s*(\d{1,2})/(\d{1,2})/(\d{4})\s*$', value) or re.match(r'^\s*(\d{4})-(\d{1,2})-(\d{1,2})\s*$', value)
        if not match:
            return None
        if len(match.group(1)) == 4:  # yyyy-mm-dd
            year, month, day = int(match.group(1)), int(match.group(2)), int(match.group(3))
        else:
            month, day, year = int(match.group(1)), int(match.group(2)), int(match.group(3))
        try:
            date = datetime(year, month, day)
            return (day, date.strftime('%B'), year)
        except ValueError:
            return None

    def format_legal_date_from_string(value: str) -> str:
        """Format for {{FORMATION_DATE}} / {{Date_of_formation_LLC}} (body of doc). Keeps existing style."""
        if not value:
            return value
        if 'day of' in value:
            match_words = re.match(r'^\s*(\d{1,2})(st|nd|rd|th)\s+day of\s+(.+)\s*$', value, re.IGNORECASE)
            if match_words:
                return value.strip()
            return value
        parsed = _parse_date(value)
        if not parsed:
            return value
        day, month_name, year = parsed
        return f"{number_to_word(day)} day of {month_name}, {year}"

    def format_witness_date_from_string(value: str) -> str:
        """Numeric ordinal only for 'IN WITNESS WHEREOF, this Resolution...' block: e.g. 8th day of February, 2026."""
        if not value:
            return value
        if 'day of' in value:
            match_words = re.match(r'^\s*(\d{1,2})(st|nd|rd|th)\s+day of\s+(.+)\s*$', value, re.IGNORECASE)
            if match_words:
                return value.strip()
            return value
        parsed = _parse_date(value)
        if not parsed:
            return value
        day, month_name, year = parsed
        return f"{day}{ordinal_suffix(day)} day of {month_name}, {year}"

    formation_date = format_legal_date_from_string(formation_date_raw)
    # Witness date should follow Payment Date (raw) like other docs
    witness_date = format_witness_date_from_string(payment_date_raw)

    # Members / managers
    members = data.get('members', []) or []
    managers = data.get('managers', []) or []
    print(f"===> Found {len(members)} members and {len(managers)} managers in form data")

    # Determine managed type:
    # - If ALL members are also ALL managers (same set), it's "member-managed"
    # - Otherwise, it's "manager-managed"
    if len(managers) == 0:
        managed_type = "member"
    elif len(members) == len(managers) and len(members) > 0:
        # Check if all members are also managers (same set)
        member_names = set(m.get('name', '').strip().lower() for m in members if m.get('name'))
        manager_names = set(m.get('name', '').strip().lower() for m in managers if m.get('name'))
        if member_names == manager_names and len(member_names) > 0:
            managed_type = "member"
        else:
            managed_type = "manager"
    else:
        managed_type = "manager"

    replacements = []
    pct_placeholders = []

    def add_placeholder(placeholder: str, value: str, bold: bool = False):
        replacements.append((placeholder, value or '', bold))

    # Generic placeholders
    add_placeholder('{{COMPANY_NAME}}', company_name, True)
    add_placeholder('{{COMPANY_ADDRESS}}', company_address)
    add_placeholder('{{FORMATION_STATE}}', formation_state)
    add_placeholder('{{FORMATION_DATE}}', formation_date)
    # Airtable-style company placeholders used in Inc minutes template
    add_placeholder('{{Company Name}}', company_name, True)
    add_placeholder('{{Company Address}}', company_address)

    # Legacy LLC-specific placeholders
    add_placeholder('{{llc_name_text}}', company_name, True)
    add_placeholder('{{full_llc_address}}', company_address)
    add_placeholder('{{full_state}}', formation_state)
    add_placeholder('{{full_state_caps}}', formation_state.upper() if formation_state else '')
    add_placeholder('{{Date_of_formation_LLC}}', formation_date)

    # Managed type placeholders
    add_placeholder('{{managed_type}}', managed_type)
    add_placeholder('{{MANAGED_TYPE}}', managed_type.upper())

    # Member placeholders: {{member_01_full_name}}, {{member_01_pct}}, etc.
    # Also alias to Airtable-style shareholder placeholders for Inc minutes template
    for idx, member in enumerate(members, start=1):
        num2 = f"{idx:02d}"
        member_name = member.get('name', '')
        for ph in (
            '{{' + f'member_{num2}_full_name' + '}}',
            '{{' + f'member_{idx}_full_name' + '}}',
            '{{' + f'Member_{num2}_full_name' + '}}',
            '{{' + f'Member_{idx}_full_name' + '}}',
        ):
            add_placeholder(ph, member_name, True)

        ownership_pct = member.get('ownershipPercent', 0)
        if ownership_pct is None:
            ownership_pct = 0
        print(f"===> Member {idx} ownership: {ownership_pct} (raw), will format to: {format_percentage(ownership_pct)}")
        pct_str = format_percentage(ownership_pct)
        pct_str_no_percent = pct_str.rstrip('%')
        for ph in (
            '{{' + f'member_{num2}_pct' + '}}',
            '{{' + f'member_{idx}_pct' + '}}',
            '{{' + f'Member_{num2}_pct' + '}}',
            '{{' + f'Member_{idx}_pct' + '}}',
        ):
            pct_placeholders.append((ph, pct_str, pct_str_no_percent))

        # Shareholder placeholders (C-Corp org minutes: shareholder instead of member)
        shares_val = member.get('shares')
        if shares_val is not None:
            # Format shares with thousands separator (e.g. 1,000)
            shares_str = f"{int(shares_val):,}"
            for ph in (
                '{{' + f'shareholder_{num2}_full_name' + '}}',
                '{{' + f'shareholder_{idx}_full_name' + '}}',
                '{{' + f'Shareholder_{num2}_full_name' + '}}',
                '{{' + f'Shareholder_{idx}_full_name' + '}}',
            ):
                add_placeholder(ph, member_name, True)
            for ph in (
                '{{' + f'shareholder_{num2}_pct' + '}}',
                '{{' + f'shareholder_{idx}_pct' + '}}',
                '{{' + f'Shareholder_{num2}_pct' + '}}',
                '{{' + f'Shareholder_{idx}_pct' + '}}',
            ):
                pct_placeholders.append((ph, pct_str, pct_str_no_percent))
            for ph in (
                '{{' + f'shareholder_{num2}_shares' + '}}',
                '{{' + f'shareholder_{idx}_shares' + '}}',
                '{{' + f'Shareholder_{num2}_shares' + '}}',
                '{{' + f'Shareholder_{idx}_shares' + '}}',
            ):
                add_placeholder(ph, shares_str, False)

        # Airtable-style placeholders for Inc/Corp 216 templates: Owner 1-6
        for ph in (
            '{{' + f'Owner {idx} Name' + '}}',
            '{{ ' + f'Owner {idx} Name' + ' }}',
        ):
            add_placeholder(ph, member_name, True)
        for ph in (
            '{{' + f'Owner {idx} Ownership %' + '}}',
            '{{ ' + f'Owner {idx} Ownership %' + ' }}',
        ):
            pct_placeholders.append((ph, pct_str, pct_str_no_percent))
        if shares_val is not None:
            shares_str = f"{int(shares_val):,}"
            for ph in (
                '{{' + f'Owner {idx} Ownership #Shares' + '}}',
                '{{ ' + f'Owner {idx} Ownership #Shares' + ' }}',
            ):
                add_placeholder(ph, shares_str, False)

    # Signature line: under signature we want "100% Owner, President and Director" (not "100% Owner, and President")
    first_manager_role_raw = ''
    first_manager_role_trimmed = ''
    signature_line_role = ''

    # Manager placeholders: {{manager_01_full_name}}, {{Manager_1}}, etc. + Officer role + Director alias
    for idx, manager in enumerate(managers, start=1):
        num2 = f"{idx:02d}"
        manager_name = manager.get('name', '')
        manager_role_raw = (manager.get('role') or '').strip() or ('President; Director' if idx == 1 else '')
        # If template already adds '; Director', trim trailing '; Director' from role to avoid 'Director; Director'
        manager_role = manager_role_raw
        if manager_role.upper().endswith('; DIRECTOR'):
            manager_role = manager_role[:-len('; Director')].rstrip()
        if idx == 1:
            first_manager_role_raw = manager_role_raw
            first_manager_role_trimmed = manager_role
            signature_line_role = manager_role_raw.replace('; ', ' and ').strip()  # "President; Director" -> "President and Director"
        for ph in (
            '{{' + f'manager_{num2}_full_name' + '}}',
            '{{' + f'manager_{idx}_full_name' + '}}',
            '{{' + f'Manager_{num2}_full_name' + '}}',
            '{{' + f'Manager_{idx}_full_name' + '}}',
            '{{' + f'Manager_{num2}' + '}}',
            '{{' + f'Manager_{idx}' + '}}',
        ):
            add_placeholder(ph, manager_name, True)
        for ph in (
            '{{' + f'Officer_{num2}_role' + '}}',
            '{{' + f'Officer_{idx}_role' + '}}',
            '{{' + f'Officer_{num2}_Role' + '}}',
            '{{' + f'Officer_{idx}_Role' + '}}',
            '{{' + f'Manager_{num2}_title' + '}}',
            '{{' + f'Manager_{idx}_title' + '}}',
        ):
            add_placeholder(ph, manager_role, False)
        # Airtable-style Officer N Name and Officer N Role placeholders (all 1-6)
        for ph in (
            '{{' + f'Officer {idx} Name' + '}}',
            '{{ ' + f'Officer {idx} Name' + ' }}',
        ):
            add_placeholder(ph, manager_name, True)
        for ph in (
            '{{' + f'Officer {idx} Role' + '}}',
            '{{ ' + f'Officer {idx} Role' + ' }}',
        ):
            add_placeholder(ph, manager_role, False)

    # Director placeholders: read from separate 'directors' array (Board of Directors)
    # Falls back to managers if no directors array is provided (backward compat)
    directors = data.get('directors', []) or []
    if not directors:
        # Fallback: use managers as directors (old behavior)
        directors = managers
    for idx, director in enumerate(directors, start=1):
        num2 = f"{idx:02d}"
        director_name = director.get('name', '')
        # Indexed placeholders: Director_1_Name, Director_01_Name
        for ph in (
            '{{' + f'Director_{num2}_Name' + '}}',
            '{{' + f'Director_{idx}_Name' + '}}',
        ):
            add_placeholder(ph, director_name, True)
        # Airtable-style placeholders: {{Director 1 Name}} through {{Director 6 Name}}
        for ph in (
            '{{' + f'Director {idx} Name' + '}}',
            '{{ ' + f'Director {idx} Name' + ' }}',
        ):
            add_placeholder(ph, director_name, True)

    def insert_run_after(paragraph, run, text):
        new_run = paragraph.add_run(text)
        run._element.addnext(new_run._element)
        return new_run

    def apply_run_format(source_run, target_run, bold_override=None):
        if source_run._element.rPr is not None:
            target_run._element.insert(0, deepcopy(source_run._element.rPr))
        if bold_override is not None:
            target_run.bold = bold_override

    def replace_span_in_paragraph(paragraph, start, end, value, bold_value=False):
        pos = 0
        run_start_idx = None
        run_end_idx = None
        run_start_offset = 0
        run_end_offset = 0

        for i, run in enumerate(paragraph.runs):
            run_len = len(run.text)
            if run_start_idx is None and pos + run_len > start:
                run_start_idx = i
                run_start_offset = start - pos
            if run_start_idx is not None and pos + run_len >= end:
                run_end_idx = i
                run_end_offset = end - pos
                break
            pos += run_len

        if run_start_idx is None or run_end_idx is None:
            return False

        if not bold_value:
            if run_start_idx == run_end_idx:
                run = paragraph.runs[run_start_idx]
                run.text = run.text[:run_start_offset] + value + run.text[run_end_offset:]
            else:
                start_run = paragraph.runs[run_start_idx]
                end_run = paragraph.runs[run_end_idx]
                start_run.text = start_run.text[:run_start_offset] + value
                for i in range(run_start_idx + 1, run_end_idx):
                    paragraph.runs[i].text = ''
                end_run.text = end_run.text[run_end_offset:]
            return True

        # Bold only the placeholder value (not the whole run)
        if run_start_idx == run_end_idx:
            run = paragraph.runs[run_start_idx]
            prefix = run.text[:run_start_offset]
            suffix = run.text[run_end_offset:]

            if prefix:
                run.text = prefix
                run.bold = False
                value_run = insert_run_after(paragraph, run, value)
                apply_run_format(run, value_run, True)
            else:
                run.text = value
                run.bold = True
                value_run = run

            if suffix:
                suffix_run = insert_run_after(paragraph, value_run, suffix)
                apply_run_format(run, suffix_run, False)
        else:
            start_run = paragraph.runs[run_start_idx]
            end_run = paragraph.runs[run_end_idx]
            prefix = start_run.text[:run_start_offset]
            suffix = end_run.text[run_end_offset:]

            for i in range(run_start_idx + 1, run_end_idx):
                paragraph.runs[i].text = ''

            if prefix:
                start_run.text = prefix
                start_run.bold = False
                value_run = insert_run_after(paragraph, start_run, value)
                apply_run_format(start_run, value_run, True)
            else:
                start_run.text = value
                start_run.bold = True
                value_run = start_run

            if suffix:
                end_run.text = suffix
                end_run.bold = False
            else:
                end_run.text = ''

        return True

    def replace_placeholder_in_paragraph(paragraph, placeholder, value, bold_value=False):
        if placeholder not in paragraph.text:
            return
        full_text = ''.join(run.text for run in paragraph.runs)
        search_start = 0
        while True:
            idx = full_text.find(placeholder, search_start)
            if idx == -1:
                break
            end = idx + len(placeholder)
            if not replace_span_in_paragraph(paragraph, idx, end, value, bold_value):
                break
            full_text = ''.join(run.text for run in paragraph.runs)
            search_start = idx + len(value)

    def replace_pct_placeholder_in_paragraph(paragraph, placeholder, pct_with_percent, pct_without_percent):
        if placeholder not in paragraph.text:
            return
        full_text = ''.join(run.text for run in paragraph.runs)
        search_start = 0
        while True:
            idx = full_text.find(placeholder, search_start)
            if idx == -1:
                break
            end = idx + len(placeholder)
            has_percent = end < len(full_text) and full_text[end] == '%'
            value = pct_without_percent if has_percent else pct_with_percent
            if not replace_span_in_paragraph(paragraph, idx, end, value):
                break
            full_text = ''.join(run.text for run in paragraph.runs)
            search_start = idx + len(value)

    def normalize_bold_for_name_paragraph(paragraph):
        text = paragraph.text
        if not text:
            return
        lower = text.lower()
        has_company = '{{COMPANY_NAME}}' in text or '{{llc_name_text}}' in text
        has_member = '{{member' in lower
        has_manager = '{{manager' in lower
        if not (has_company or has_member or has_manager):
            return
        # Only placeholders + "RESOLVED," should be bold; reset others in this paragraph
        for run in paragraph.runs:
            run.bold = False
        for run in paragraph.runs:
            if 'RESOLVED,' in run.text:
                run.bold = True

    def replace_date_in_paragraph(paragraph, date_value):
        if "IN WITNESS WHEREOF" not in paragraph.text:
            return
        full_text = ''.join(run.text for run in paragraph.runs)
        # Replace mm/dd/yyyy (e.g. 02/08/2026)
        match = re.search(r'\b\d{2}/\d{2}/\d{4}\b', full_text)
        if match:
            replace_span_in_paragraph(paragraph, match.start(), match.end(), date_value, False)
            return
        # Replace "this [ordinal word] day of [Month], [year]" (LLC/Corp templates)
        match = re.search(
            r'\bthis\s+(?:first|second|third|fourth|fifth|sixth|seventh|eighth|ninth|tenth|'
            r'eleventh|twelfth|thirteenth|fourteenth|fifteenth|sixteenth|seventeenth|eighteenth|'
            r'nineteenth|twentieth|twenty-first|twenty-second|twenty-third|twenty-fourth|'
            r'twenty-fifth|twenty-sixth|twenty-seventh|twenty-eighth|twenty-ninth|thirtieth|'
            r'thirty-first)\s+day of\s+[A-Za-z]+\s*,\s*\d{4}\b',
            full_text,
        )
        if match:
            replace_span_in_paragraph(paragraph, match.start(), match.end(), date_value, False)
            return
        # Replace numeric ordinal form: "this 10th day of July, 2025" (Inc minutes template)
        match = re.search(
            r'\bthis\s+\d{1,2}(st|nd|rd|th)\s+day of\s+[A-Za-z]+\s*,\s*\d{4}\b',
            full_text,
        )
        if match:
            replacement = f"this {date_value}"
            replace_span_in_paragraph(paragraph, match.start(), match.end(), replacement, False)

    # Use numeric ordinal in IN WITNESS WHEREOF block for both LLC and C-Corp/S-Corp templates
    WITNESS_DATE_PLACEHOLDERS = ('{{FORMATION_DATE}}', '{{Date_of_formation_LLC}}')

    def replace_all_in_paragraph(paragraph):
        if '{{' not in paragraph.text:
            replace_date_in_paragraph(paragraph, witness_date)
            return
        replace_date_in_paragraph(paragraph, witness_date)
        normalize_bold_for_name_paragraph(paragraph)
        in_witness_block = "IN WITNESS WHEREOF" in paragraph.text
        for placeholder, value, bold_value in replacements:
            # In witness block use numeric ordinal (8th day of...) for date placeholders
            use_value = witness_date if (in_witness_block and placeholder in WITNESS_DATE_PLACEHOLDERS) else value
            replace_placeholder_in_paragraph(paragraph, placeholder, use_value, bold_value)
        for placeholder, pct_str, pct_str_no_percent in pct_placeholders:
            replace_pct_placeholder_in_paragraph(paragraph, placeholder, pct_str, pct_str_no_percent)

    # Replace in paragraphs (preserve formatting by editing runs in place)
    for paragraph in doc.paragraphs:
        replace_all_in_paragraph(paragraph)

    # Replace in tables (cell paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_all_in_paragraph(paragraph)

    # Fix signature-block indentation: normalize ALL lines in the signature
    # section to use exactly 6 leading tabs.  Some LLC templates (3+ members
    # with 2+ managers) use SPACES for indentation on member 3+ blocks.
    # If we blindly prepend tabs without stripping spaces first, the line
    # overflows and wraps.  This code normalizes every meaningful signature
    # line (By:, Name:, Title:, and "% Owner") to 6-tab indentation.
    SIG_LINE_TAB_COUNT = '\t\t\t\t\t\t'  # 6 tabs — matches template convention
    in_sig = False
    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        if 'WITNESS WHEREOF' in txt:
            in_sig = True
        if not in_sig or not paragraph.runs:
            continue
        first_run = paragraph.runs[0]
        # Detect signature lines that need tab normalization:
        #   "XX% Owner …" — ownership percentage
        #   "By: ___" — signature line
        #   "Name: …" — member/shareholder name
        #   "Title: …" — title line
        needs_fix = False
        if re.match(r'^\d+%\s*Owner', txt):
            needs_fix = True
        elif re.match(r'^By:\s*_', txt):
            needs_fix = True
        elif re.match(r'^Name:\s', txt):
            needs_fix = True
        elif re.match(r'^Title:\s', txt):
            needs_fix = True
        if needs_fix:
            raw = first_run.text
            stripped = raw.lstrip()
            if stripped:
                # First run has the content — strip whitespace and add tabs
                if not raw.startswith(SIG_LINE_TAB_COUNT):
                    first_run.text = SIG_LINE_TAB_COUNT + stripped
            else:
                # First run is only whitespace (tabs/spaces); normalize to 6 tabs
                first_run.text = SIG_LINE_TAB_COUNT

    # --- #7: Dynamic authority clause ---
    # For companies with >1 officer or >1 director, replace "the President" with
    # "any officer or director" in banking/authority RESOLVED clauses.
    num_officers = len(managers) if managers else 1
    num_directors = len(directors) if directors else 1
    if num_officers > 1 or num_directors > 1:
        authority_replacements = [
            ("the President be and hereby is authorized to open a bank account",
             "any officer or director of the Company be and hereby is authorized to open a bank account"),
            ("the Company's President is authorized to execute",
             "any officer or director of the Company is authorized to execute"),
            ("as the President determines",
             "as the officers or directors determine"),
            ("actions taken by the President of the Company",
             "actions taken by any officer or director of the Company"),
            ("designated by her",
             "designated by them"),
            ("designated by him",
             "designated by them"),
        ]
        for paragraph in doc.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            for old_text, new_text in authority_replacements:
                if old_text in full_text:
                    idx = full_text.find(old_text)
                    replace_span_in_paragraph(paragraph, idx, idx + len(old_text), new_text, False)
                    full_text = ''.join(run.text for run in paragraph.runs)

    # Fix signature line: "100% Owner, and President" -> "100% Owner, President and Director"
    if first_manager_role_trimmed and signature_line_role and first_manager_role_trimmed != signature_line_role:
        old_suffix = " and " + first_manager_role_trimmed
        new_suffix = " " + signature_line_role
        for paragraph in doc.paragraphs:
            if old_suffix in paragraph.text:
                full_text = ''.join(run.text for run in paragraph.runs)
                idx = full_text.find(old_suffix)
                if idx != -1:
                    replace_span_in_paragraph(paragraph, idx, idx + len(old_suffix), new_suffix, False)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_suffix in paragraph.text:
                            full_text = ''.join(run.text for run in paragraph.runs)
                            idx = full_text.find(old_suffix)
                            if idx != -1:
                                replace_span_in_paragraph(paragraph, idx, idx + len(old_suffix), new_suffix, False)

    # Ensure heading (first paragraph) is all caps company name like template style
    if doc.paragraphs:
        heading_para = doc.paragraphs[0]
        full_text = ''.join(run.text for run in heading_para.runs) or heading_para.text
        if full_text.strip():
            upper = full_text.upper()
            if heading_para.runs:
                heading_para.runs[0].text = upper
                for r in heading_para.runs[1:]:
                    r.text = ''
            else:
                heading_para.text = upper

    print("===> Placeholders replaced successfully")


# ---------- Post-processing: best-practice formatting fixes ----------

def _add_keep_next(paragraph):
    """Add keepNext property to a paragraph so it stays with the next paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:keepNext')):
        pPr.remove(existing)
    keep_next = OxmlElement('w:keepNext')
    pPr.append(keep_next)


def post_process_org_resolution(doc):
    """Best-practice formatting fixes for Org Resolution documents:
    1. Ensure [SIGNATURE PAGE BELOW] + pageBreakBefore on IN WITNESS WHEREOF
    2. Remove excess empty paragraphs in signature section (keep max 2)
    3. Remove "PAGE X" footer text
    4. keepNext on RESOLVED headings so they don't orphan at page bottom
    """
    print("===> Post-processing org resolution...")

    paragraphs = doc.paragraphs

    # --- 1. Ensure [SIGNATURE PAGE BELOW] + pageBreakBefore ---
    # The 216 templates already have this, but handle any template that might not.
    witness_para = None
    for paragraph in paragraphs:
        if 'IN WITNESS WHEREOF' in paragraph.text:
            witness_para = paragraph
            break

    if witness_para is not None:
        body = witness_para._p.getparent()

        # Check if [SIGNATURE PAGE BELOW] already exists
        prev = witness_para._p.getprevious()
        prev_text = ''
        if prev is not None:
            prev_txts = []
            for r_el in prev.iterchildren(qn('w:r')):
                for t_el in r_el.iterchildren(qn('w:t')):
                    prev_txts.append(t_el.text or '')
            prev_text = ''.join(prev_txts).strip()

        if 'SIGNATURE PAGE BELOW' not in prev_text:
            # Remove trailing empties before witness
            while True:
                prev_el = witness_para._p.getprevious()
                if prev_el is None:
                    break
                prev_txts = []
                for r_el in prev_el.iterchildren(qn('w:r')):
                    for t_el in r_el.iterchildren(qn('w:t')):
                        prev_txts.append(t_el.text or '')
                if ''.join(prev_txts).strip() == '':
                    body.remove(prev_el)
                else:
                    break

            # Insert [SIGNATURE PAGE BELOW] centered
            sig_p = OxmlElement('w:p')
            sig_pPr = OxmlElement('w:pPr')
            sig_jc = OxmlElement('w:jc')
            sig_jc.set(qn('w:val'), 'center')
            sig_pPr.append(sig_jc)
            sig_p.append(sig_pPr)
            sig_r = OxmlElement('w:r')
            sig_rPr = OxmlElement('w:rPr')
            sig_rFonts = OxmlElement('w:rFonts')
            sig_rFonts.set(qn('w:ascii'), 'Times New Roman')
            sig_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            sig_rPr.append(sig_rFonts)
            sig_sz = OxmlElement('w:sz')
            sig_sz.set(qn('w:val'), '24')
            sig_rPr.append(sig_sz)
            sig_r.append(sig_rPr)
            sig_t = OxmlElement('w:t')
            sig_t.text = '[SIGNATURE PAGE BELOW]'
            sig_r.append(sig_t)
            sig_p.append(sig_r)
            body.insert(list(body).index(witness_para._p), sig_p)
            print("===> Inserted [SIGNATURE PAGE BELOW]")

        # Ensure pageBreakBefore on WITNESS WHEREOF
        pPr = witness_para._p.get_or_add_pPr()
        if not pPr.findall(qn('w:pageBreakBefore')):
            page_break = OxmlElement('w:pageBreakBefore')
            pPr.append(page_break)
            print("===> Added pageBreakBefore to IN WITNESS WHEREOF")

    # --- 2. Remove excess empty paragraphs in signature section (keep max 2) ---
    in_sig = False
    to_remove = []
    consecutive = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if 'IN WITNESS WHEREOF' in text:
            in_sig = True
            consecutive = 0
            continue
        if in_sig:
            if not text:
                consecutive += 1
                if consecutive > 2:
                    to_remove.append(paragraph)
            else:
                consecutive = 0
    removed = 0
    for p in to_remove:
        parent = p._p.getparent()
        if parent is not None:
            parent.remove(p._p)
            removed += 1

    # --- 3. Remove "PAGE X" footer text ---
    pages_removed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r'^PAGE\s+\d+$', text):
            parent = paragraph._p.getparent()
            if parent is not None:
                parent.remove(paragraph._p)
                pages_removed += 1

    # --- 4. keepNext on RESOLVED paragraphs ---
    # Ensure "RESOLVED," paragraphs stay with their body text
    resolved_fixed = 0
    all_paras = doc.paragraphs
    for i, paragraph in enumerate(all_paras):
        text = paragraph.text.strip()
        if text.upper().startswith('RESOLVED') and len(text) < 200:
            _add_keep_next(paragraph)
            resolved_fixed += 1
            # Also keepNext on following empty para if any
            if i + 1 < len(all_paras) and not all_paras[i + 1].text.strip():
                _add_keep_next(all_paras[i + 1])

    print(f"===> Post-processing done: {removed} excess empties removed, {pages_removed} PAGE X removed, {resolved_fixed} RESOLVED keepNext")


def lambda_handler(event, context):
    print("===> RAW EVENT:")
    print(json.dumps(event))
    
    try:
        # Parse request body
        if "body" in event:
            body = json.loads(event["body"])
        else:
            body = event
        
        form_data = body.get("form_data")
        s3_bucket = body.get("s3_bucket")
        s3_key = body.get("s3_key")
        templateUrl = body.get("templateUrl")
        return_docx = body.get("return_docx", False)
        
        if not form_data:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Missing 'form_data'"})
            }
        
        if not s3_bucket or not s3_key:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Missing 's3_bucket' or 's3_key'"})
            }
        
        if not templateUrl:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Missing 'templateUrl'"})
            }
        
        # Extract template bucket and key from URL
        template_bucket, template_key = extract_s3_info(templateUrl)
        if not template_bucket or not template_key:
            template_bucket = TEMPLATE_BUCKET
            template_key = 'llc-formation-templates/organizational-resolution-all-templates/Template Organization Resolution_1 Member/Template Organization Resolution_1 Member_1 Manager.docx'
            print(f"===> Could not parse templateUrl, using default: s3://{template_bucket}/{template_key}")
        else:
            print(f"===> Template: s3://{template_bucket}/{template_key}")
        
        print(f"===> Output: s3://{s3_bucket}/{s3_key}")
        
    except Exception as e:
        return {
            "statusCode": 400,
            "body": json.dumps({"error": "Invalid input payload", "details": str(e)})
        }
    
    # Prepare file paths
    tmpdir = tempfile.gettempdir()
    template_path = os.path.join(tmpdir, "template_org_resolution.docx")
    output_path = os.path.join(tmpdir, "filled_org_resolution.docx")
    
    try:
        # Download template from S3
        download_from_s3(template_bucket, template_key, template_path)
        
        # Load Word document
        print("===> Loading Word document...")
        doc = Document(template_path)
        
        # Replace placeholders with form data
        replace_placeholders(doc, form_data)

        # Apply best-practice formatting fixes
        post_process_org_resolution(doc)

        # Save filled document
        print("===> Saving filled document...")
        doc.save(output_path)
        print(f"===> Saved to {output_path}")
        
        # Upload to S3
        upload_to_s3(output_path, s3_bucket, s3_key)
        
        # Return response
        if return_docx:
            with open(output_path, 'rb') as f:
                docx_content = f.read()
            
            encoded = base64.b64encode(docx_content).decode('utf-8')
            
            return {
                "statusCode": 200,
                "headers": {
                    "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "Content-Disposition": f'attachment; filename="{os.path.basename(s3_key)}"'
                },
                "body": encoded,
                "isBase64Encoded": True
            }
        else:
            return {
                "statusCode": 200,
                "body": json.dumps({
                    "message": "✅ Document uploaded to S3",
                    "s3_bucket": s3_bucket,
                    "s3_key": s3_key,
                    "s3_url": f"s3://{s3_bucket}/{s3_key}"
                })
            }
    
    except Exception as e:
        import traceback
        error_msg = str(e)
        traceback.print_exc()
        print(f"===> ERROR: {error_msg}")
        return {
            "statusCode": 500,
            "body": json.dumps({
                "error": "Failed to generate organizational resolution",
                "details": error_msg
            })
        }
