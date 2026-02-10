#!/usr/bin/env python3
"""
Create C-Corp Organizational Minutes template from Gym Kidz source and upload to S3.
Replaces company-specific text with placeholders the organizational-resolution Lambda expects.
Usage:
  python scripts/create-organizational-minutes-template.py [path_to_source.docx]
Default source: organizational-resolution-inc/organizational-minutes-gym-kidz-SOURCE.docx
Requires: python-docx, boto3 (pip install python-docx boto3)
Env: AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, AWS_REGION (default us-west-1)
     TEMPLATE_BUCKET (default company-formation-template-llc-and-inc)
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Install python-docx: pip install python-docx")
    sys.exit(1)

# Replacements: (search, placeholder) — order matters (longer strings first)
# IMPORTANT: We only replace inside existing runs (no paragraph-level rewrites)
# to avoid changing spacing/formatting. This keeps the original Gym Kidz
# layout intact while swapping text for placeholders.
REPLACEMENTS = [
    # Company name (multiple variants)
    ("GYM KIDZ IN MIAMI LAKES INC.", "{{COMPANY_NAME}}"),
    ("Gym Kidz in Miami Lakes Inc.", "{{COMPANY_NAME}}"),
    ("Gym Kidz in Miami Lakes Inc", "{{COMPANY_NAME}}"),
    ("GYM KIDZ IN MIAMI LAKES INC", "{{COMPANY_NAME}}"),
    # Company address
    ("2038 NE 155 Street, North Miami Beach, FL 33162", "{{COMPANY_ADDRESS}}"),
    # State (only when referring to formation state)
    ("State of Florida", "State of {{FORMATION_STATE}}"),
    ("the State of Florida", "the State of {{FORMATION_STATE}}"),
    ("Florida corporation", "{{FORMATION_STATE}} corporation"),
    ("Florida Secretary of State", "{{FORMATION_STATE}} Secretary of State"),
    # Shareholder name (C-Corp: shareholder_1, maps from Airtable Owner 1 Name)
    ("Joslyn Varona", "{{shareholder_1_full_name}}"),
    # Shares and % (base 1,000 shares; Lambda fills shareholder_1_shares and shareholder_1_pct from Owner Ownership %)
    ("1,000 Shares, or 100%", "{{shareholder_1_shares}} Shares, or {{shareholder_1_pct}}"),
    ("100% of the Company", "{{shareholder_1_pct}} of the Company"),
    ("100% Owner", "{{shareholder_1_pct}} Owner"),
    # Officer title (Airtable Officer 1 Role → President; Director or custom)
    ("President; Director", "{{Officer_1_role}}"),
    # IN WITNESS date: Lambda replaces date in paragraph containing "IN WITNESS WHEREOF"
    ("____ day of ____________, 20___", "{{FORMATION_DATE}}"),
]


def replace_in_paragraph(paragraph):
    """Replace all replacement pairs in each run (no structural changes)."""
    if not paragraph.text.strip():
        return
    for run in paragraph.runs:
        if not run.text:
            continue
        text = run.text
        changed = False
        for search, placeholder in REPLACEMENTS:
            if search in text:
                text = text.replace(search, placeholder)
                changed = True
        if changed:
            run.text = text


def replace_in_doc(doc):
    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)


def main():
    repo_root = Path(__file__).resolve().parent.parent
    default_source = repo_root / "organizational-resolution-inc" / "organizational-minutes-gym-kidz-SOURCE.docx"
    source_path = Path(sys.argv[1]) if len(sys.argv) > 1 else default_source
    if not source_path.exists():
        print(f"Source not found: {source_path}")
        print("Usage: python create-organizational-minutes-template.py [path_to_source.docx]")
        sys.exit(1)
    out_dir = repo_root / "organizational-resolution-inc"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "organizational-minutes-inc-1.docx"
    print(f"Loading {source_path}...")
    doc = Document(str(source_path))
    print("Applying placeholder replacements...")
    replace_in_doc(doc)
    print(f"Saving {out_path}...")
    doc.save(str(out_path))
    print(f"Template saved: {out_path}")

    # Upload to S3
    bucket = os.environ.get("TEMPLATE_BUCKET", "company-formation-template-llc-and-inc")
    region = os.environ.get("AWS_REGION", "us-west-1")
    s3_key = "templates/organizational-resolution-inc/organizational-minutes-inc-1.docx"
    try:
        import boto3
        # Use default credential chain (env vars, ~/.aws/credentials, AWS_PROFILE)
        s3 = boto3.client("s3", region_name=region)
        s3.upload_file(str(out_path), bucket, s3_key, ExtraArgs={"ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})
        print(f"Uploaded to s3://{bucket}/{s3_key}")
        # Copy to -2 through -6 so API can pick by member count
        for n in range(2, 7):
            copy_key = f"templates/organizational-resolution-inc/organizational-minutes-inc-{n}.docx"
            s3.copy_object(Bucket=bucket, CopySource={"Bucket": bucket, "Key": s3_key}, Key=copy_key)
            print(f"Copied to {copy_key}")
    except ImportError:
        print("Install boto3 to upload: pip install boto3. Skipping upload.")
    except Exception as e:
        print(f"Upload failed: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
